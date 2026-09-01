"""
Shaxsiy AI Yordamchi — Telegram Bot (v2)
=========================================
Imkoniyatlar:
- Matn, GOLOS va RASM xabarlarni tushunadi (Gemini)
- Internetdan yangi ma'lumot qidiradi (Google Search)
- Shaxsiy buxgalter: kirim-chiqim + hisobot (chek rasmidan ham o'qiydi!)
- Eslatmalar: "Ertaga 9 da dorini eslatib qo'y" — vaqtida xabar keladi
- Qaydlar: "Eslab qol: ..." — keyin so'rasangiz topib beradi

Ishga tushirish:
  1. .env faylga BOT_TOKEN va GEMINI_API_KEY yozing
  2. pip install -r requirements.txt
  3. python bot.py
"""

import asyncio
import io
import logging
import os
import time
from datetime import datetime, timedelta
from zoneinfo import ZoneInfo

import requests
import psycopg2
from psycopg2.extras import RealDictCursor

from aiogram import Bot, Dispatcher, F, Router
from aiogram.enums import ChatAction
from aiogram.filters import Command, CommandStart
from aiogram.types import (
    Message, InlineKeyboardMarkup, InlineKeyboardButton, CallbackQuery, BufferedInputFile,
    BotCommand,
)
from apscheduler.schedulers.asyncio import AsyncIOScheduler
from dotenv import load_dotenv
from google import genai
from google.genai import types

load_dotenv()

BOT_TOKEN = os.getenv("BOT_TOKEN")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
MODEL = os.getenv("GEMINI_MODEL", "gemini-3.5-flash")
DATABASE_URL = os.getenv("DATABASE_URL")
ADMIN_ID = int(os.getenv("ADMIN_ID", "0"))
GOOGLE_CSE_KEY = os.getenv("GOOGLE_CSE_KEY", "")
GOOGLE_CSE_ID = os.getenv("GOOGLE_CSE_ID", "")
TZ = ZoneInfo(os.getenv("TZ", "Asia/Tashkent"))

bot_enabled = True  # admin o'chirishi mumkin

if not BOT_TOKEN or not GEMINI_API_KEY:
    raise SystemExit("Xato: .env faylda BOT_TOKEN va GEMINI_API_KEY bo'lishi shart!")
if not DATABASE_URL:
    raise SystemExit("Xato: DATABASE_URL bo'lishi shart!")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("ai-assistant")

client = genai.Client(api_key=GEMINI_API_KEY)
scheduler = AsyncIOScheduler(timezone=TZ)
BOT: Bot | None = None  # main() da to'ldiriladi


def now_local() -> datetime:
    return datetime.now(TZ)


# ============================================================
# BAZA (PostgreSQL)
# ============================================================

def get_conn():
    return psycopg2.connect(DATABASE_URL, cursor_factory=RealDictCursor)


def db_init():
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                CREATE TABLE IF NOT EXISTS transactions (
                    id SERIAL PRIMARY KEY,
                    user_id BIGINT NOT NULL,
                    type TEXT NOT NULL CHECK(type IN ('kirim', 'chiqim')),
                    amount REAL NOT NULL,
                    category TEXT NOT NULL,
                    note TEXT,
                    created_at TIMESTAMP NOT NULL DEFAULT NOW()
                )
            """)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS reminders (
                    id SERIAL PRIMARY KEY,
                    user_id BIGINT NOT NULL,
                    text TEXT NOT NULL,
                    remind_at TEXT NOT NULL,
                    sent INTEGER NOT NULL DEFAULT 0
                )
            """)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS notes (
                    id SERIAL PRIMARY KEY,
                    user_id BIGINT NOT NULL,
                    text TEXT NOT NULL,
                    created_at TIMESTAMP NOT NULL DEFAULT NOW()
                )
            """)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS users (
                    user_id BIGINT PRIMARY KEY,
                    username TEXT,
                    full_name TEXT,
                    first_seen TIMESTAMP NOT NULL DEFAULT NOW(),
                    last_seen TIMESTAMP NOT NULL DEFAULT NOW(),
                    message_count INTEGER NOT NULL DEFAULT 0,
                    approved BOOLEAN NOT NULL DEFAULT FALSE
                )
            """)
            cur.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS approved BOOLEAN NOT NULL DEFAULT FALSE")
            # Snippet kutubxonasi: qaydlarga sarlavha, teglar, kod bloki (additiv)
            cur.execute("ALTER TABLE notes ADD COLUMN IF NOT EXISTS title TEXT NOT NULL DEFAULT ''")
            cur.execute("ALTER TABLE notes ADD COLUMN IF NOT EXISTS tags TEXT NOT NULL DEFAULT ''")
            cur.execute("ALTER TABLE notes ADD COLUMN IF NOT EXISTS code TEXT NOT NULL DEFAULT ''")
            cur.execute("ALTER TABLE notes ADD COLUMN IF NOT EXISTS lang TEXT NOT NULL DEFAULT ''")
            cur.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS reply_mode TEXT NOT NULL DEFAULT 'text'")
            cur.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS business_auto BOOLEAN NOT NULL DEFAULT FALSE")
            cur.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS business_info TEXT NOT NULL DEFAULT ''")
            cur.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS business_hours TEXT NOT NULL DEFAULT ''")
            cur.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS business_notify BOOLEAN NOT NULL DEFAULT TRUE")
            cur.execute("""
                CREATE TABLE IF NOT EXISTS business_chats (
                    owner_id BIGINT NOT NULL,
                    chat_id BIGINT NOT NULL,
                    conn_id TEXT NOT NULL,
                    name TEXT,
                    username TEXT,
                    last_seen TIMESTAMP NOT NULL DEFAULT NOW(),
                    PRIMARY KEY (owner_id, chat_id)
                )
            """)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS business_log (
                    id SERIAL PRIMARY KEY,
                    owner_id BIGINT NOT NULL,
                    chat_id BIGINT NOT NULL,
                    sender TEXT,
                    text TEXT,
                    created_at TIMESTAMP NOT NULL DEFAULT NOW()
                )
            """)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS profiles (
                    user_id BIGINT PRIMARY KEY,
                    name TEXT,
                    profession TEXT,
                    interests TEXT,
                    language TEXT DEFAULT 'uz',
                    goals TEXT,
                    onboarded BOOLEAN NOT NULL DEFAULT FALSE
                )
            """)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS long_memory (
                    id SERIAL PRIMARY KEY,
                    user_id BIGINT NOT NULL,
                    summary TEXT NOT NULL,
                    created_at TIMESTAMP NOT NULL DEFAULT NOW()
                )
            """)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS chat_history (
                    id SERIAL PRIMARY KEY,
                    user_id BIGINT NOT NULL,
                    role TEXT NOT NULL,
                    text TEXT NOT NULL,
                    created_at TIMESTAMP NOT NULL DEFAULT NOW()
                )
            """)
            cur.execute(
                "CREATE INDEX IF NOT EXISTS idx_chat_history_user ON chat_history(user_id, id)"
            )
            # Ilovadagi "chatlar menyusi" uchun: har suhbat alohida kontekst (Telegram = 0)
            cur.execute(
                "ALTER TABLE chat_history ADD COLUMN IF NOT EXISTS chat_id BIGINT NOT NULL DEFAULT 0"
            )


# --- Buxgalteriya ---

def db_add_transaction(user_id: int, tx_type: str, amount: float, category: str, note: str) -> str:
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO transactions (user_id, type, amount, category, note) VALUES (%s, %s, %s, %s, %s)",
                (user_id, tx_type, amount, category, note),
            )
    return f"Yozildi: {tx_type} {amount:,.0f} so'm, kategoriya: {category}" + (f" ({note})" if note else "")


def db_get_report(user_id: int, period: str = "oy", start_date: str = "", end_date: str = "") -> str:
    now = now_local().replace(tzinfo=None)
    label = period

    # Aniq sana oralig'i berilgan bo'lsa — o'shani ishlatamiz
    if start_date:
        try:
            start = datetime.strptime(start_date, "%Y-%m-%d")
        except ValueError:
            return "Boshlanish sanasi noto'g'ri (YYYY-MM-DD bo'lishi kerak)."
        if end_date:
            try:
                end = datetime.strptime(end_date, "%Y-%m-%d").replace(hour=23, minute=59, second=59)
            except ValueError:
                return "Tugash sanasi noto'g'ri (YYYY-MM-DD bo'lishi kerak)."
        else:
            end = start.replace(hour=23, minute=59, second=59)
        label = f"{start_date}" + (f" — {end_date}" if end_date and end_date != start_date else "")
    else:
        end = now
        if period == "bugun":
            start = now.replace(hour=0, minute=0, second=0)
        elif period == "kecha":
            y = now - timedelta(days=1)
            start = y.replace(hour=0, minute=0, second=0)
            end = y.replace(hour=23, minute=59, second=59)
            label = "kecha"
        elif period == "hafta":
            start = now - timedelta(days=7)
        else:
            start = now - timedelta(days=30)

    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """SELECT type, category, SUM(amount), COUNT(*)
                   FROM transactions
                   WHERE user_id = %s AND created_at >= %s AND created_at <= %s
                   GROUP BY type, category
                   ORDER BY type, SUM(amount) DESC""",
                (user_id, start, end),
            )
            rows = cur.fetchall()

    if not rows:
        return f"Bu davr ({label}) uchun yozuvlar topilmadi."

    kirim_total, chiqim_total = 0.0, 0.0
    lines = [f"Hisobot ({label}):"]
    for row in rows:
        tx_type, category, total, count = row["type"], row["category"], row["sum"], row["count"]
        lines.append(f"- {tx_type} | {category}: {total:,.0f} so'm ({count} ta)")
        if tx_type == "kirim":
            kirim_total += total
        else:
            chiqim_total += total
    lines.append(f"Jami kirim: {kirim_total:,.0f} so'm")
    lines.append(f"Jami chiqim: {chiqim_total:,.0f} so'm")
    lines.append(f"Balans: {kirim_total - chiqim_total:,.0f} so'm")
    return "\n".join(lines)


def db_delete_last(user_id: int) -> str:
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT id, type, amount, category FROM transactions WHERE user_id = %s ORDER BY id DESC LIMIT 1",
                (user_id,),
            )
            row = cur.fetchone()
            if not row:
                return "O'chiradigan yozuv yo'q."
            cur.execute("DELETE FROM transactions WHERE id = %s", (row["id"],))
    return f"O'chirildi: {row['type']} {row['amount']:,.0f} so'm ({row['category']})"


# --- Eslatmalar ---

async def fire_reminder(reminder_id: int, user_id: int, text: str):
    """Vaqti kelganda foydalanuvchiga xabar yuboradi."""
    try:
        if BOT:
            await BOT.send_message(user_id, f"⏰ Eslatma: {text}")
        with get_conn() as conn:
            with conn.cursor() as cur:
                cur.execute("UPDATE reminders SET sent = 1 WHERE id = %s", (reminder_id,))
    except Exception:
        logger.exception("Eslatma yuborishda xato")


def schedule_reminder(reminder_id: int, user_id: int, text: str, remind_at: datetime):
    scheduler.add_job(
        fire_reminder, "date", run_date=remind_at,
        args=[reminder_id, user_id, text],
        id=f"rem_{reminder_id}", replace_existing=True,
    )


def db_set_reminder(user_id: int, text: str, remind_at_str: str) -> str:
    try:
        remind_at = datetime.strptime(remind_at_str, "%Y-%m-%d %H:%M").replace(tzinfo=TZ)
    except ValueError:
        return "Vaqt formati noto'g'ri. 'YYYY-MM-DD HH:MM' formatida bo'lishi kerak."

    if remind_at <= now_local():
        return "Bu vaqt o'tib ketgan. Kelajakdagi vaqtni ayting."

    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO reminders (user_id, text, remind_at) VALUES (%s, %s, %s) RETURNING id",
                (user_id, text, remind_at.strftime("%Y-%m-%d %H:%M")),
            )
            reminder_id = cur.fetchone()["id"]

    schedule_reminder(reminder_id, user_id, text, remind_at)
    return f"Eslatma o'rnatildi: \"{text}\" — {remind_at.strftime('%d.%m.%Y soat %H:%M')} (№{reminder_id})"


def db_list_reminders(user_id: int) -> str:
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT id, text, remind_at FROM reminders WHERE user_id = %s AND sent = 0 ORDER BY remind_at",
                (user_id,),
            )
            rows = cur.fetchall()
    if not rows:
        return "Faol eslatmalar yo'q."
    return "Faol eslatmalar:\n" + "\n".join(f"№{r['id']}: {r['text']} — {r['remind_at']}" for r in rows)


def db_delete_reminder(user_id: int, reminder_id: int) -> str:
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "DELETE FROM reminders WHERE id = %s AND user_id = %s AND sent = 0",
                (reminder_id, user_id),
            )
            deleted = cur.rowcount
    if deleted == 0:
        return f"№{reminder_id} eslatma topilmadi."
    try:
        scheduler.remove_job(f"rem_{reminder_id}")
    except Exception:
        pass
    return f"№{reminder_id} eslatma o'chirildi."


def restore_reminders():
    """Server qayta yonsa — bazadagi eslatmalarni qayta yuklaymiz."""
    now = now_local()
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT id, user_id, text, remind_at FROM reminders WHERE sent = 0")
            rows = cur.fetchall()
    restored = 0
    for row in rows:
        remind_at = datetime.strptime(row["remind_at"], "%Y-%m-%d %H:%M").replace(tzinfo=TZ)
        if remind_at <= now:
            remind_at = now + timedelta(seconds=10)
        schedule_reminder(row["id"], row["user_id"], row["text"], remind_at)
        restored += 1
    if restored:
        logger.info("%d ta eslatma qayta yuklandi", restored)


# --- Qaydlar ---

def _norm_tags(tags: str) -> str:
    """'Railway, ffmpeg' -> 'railway,ffmpeg' (kichik harf, bo'shliqsiz, # belgisiz)."""
    parts = [t.strip().lstrip("#").lower() for t in (tags or "").replace(";", ",").split(",")]
    return ",".join(sorted({p for p in parts if p}))


def db_add_note(user_id: int, text: str, title: str = "", tags: str = "", code: str = "", lang: str = "") -> str:
    tags = _norm_tags(tags)
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO notes (user_id, text, title, tags, code, lang) VALUES (%s, %s, %s, %s, %s, %s) RETURNING id",
                (user_id, text or "", title[:120], tags, code, lang.lower()[:20]),
            )
            note_id = cur.fetchone()["id"]
    label = title or (text[:60] if text else "kod snippeti")
    tag_str = f" [{', '.join('#' + t for t in tags.split(','))}]" if tags else ""
    kind = "Snippet" if code else "Qayd"
    return f"{kind} saqlandi №{note_id}: {label}{tag_str}"


def _format_note(r: dict, full: bool = True) -> str:
    head = f"№{r['id']}"
    if r.get("title"):
        head += f" — {r['title']}"
    tags = r.get("tags") or ""
    if tags:
        head += "  " + " ".join("#" + t for t in tags.split(","))
    head += f"  ({str(r['created_at'])[:10]})"
    out = head
    if r.get("text"):
        out += f"\n{r['text']}"
    if r.get("code") and full:
        out += f"\n```{r.get('lang') or ''}\n{r['code']}\n```"
    elif r.get("code"):
        first = r["code"].strip().splitlines()[0][:80] if r["code"].strip() else ""
        out += f"\n<kod: {first}...>"
    return out


def db_find_notes(user_id: int, query: str = "", tag: str = "") -> str:
    tag = _norm_tags(tag).split(",")[0] if tag else ""
    with get_conn() as conn:
        with conn.cursor() as cur:
            conds, params = ["user_id = %s"], [user_id]
            if query:
                conds.append("(text ILIKE %s OR title ILIKE %s OR code ILIKE %s OR tags ILIKE %s)")
                params += [f"%{query}%"] * 4
            if tag:
                conds.append("(',' || tags || ',') ILIKE %s")
                params.append(f"%,{tag},%")
            cur.execute(
                f"SELECT id, text, title, tags, code, lang, created_at FROM notes "
                f"WHERE {' AND '.join(conds)} ORDER BY id DESC LIMIT 10",
                params,
            )
            rows = cur.fetchall()
    if not rows:
        return "Qaydlar topilmadi." if (query or tag) else "Hali qaydlar yo'q."
    # 1-2 ta natija bo'lsa to'liq (kod bilan), ko'p bo'lsa qisqa ro'yxat
    full = len(rows) <= 2
    body = "\n\n".join(_format_note(r, full=full) for r in rows)
    hint = "" if full else "\n\nTo'liq ko'rish: qayd raqamini ayting (masalan «№12 ni ko'rsat»)."
    return f"Topilgan qaydlar ({len(rows)}):\n\n{body}{hint}"


def db_get_note(user_id: int, note_id: int) -> str:
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT id, text, title, tags, code, lang, created_at FROM notes WHERE id = %s AND user_id = %s",
                (note_id, user_id),
            )
            r = cur.fetchone()
    return _format_note(r, full=True) if r else f"№{note_id} qayd topilmadi."


def db_list_tags(user_id: int) -> str:
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT tags FROM notes WHERE user_id = %s AND tags <> ''", (user_id,))
            rows = cur.fetchall()
    counts: dict[str, int] = {}
    for r in rows:
        for t in r["tags"].split(","):
            counts[t] = counts.get(t, 0) + 1
    if not counts:
        return "Hali teglar yo'q. Qayd saqlaganda teg qo'shing: «eslab qol #railway ...»."
    top = sorted(counts.items(), key=lambda x: (-x[1], x[0]))
    return "Teglar:\n" + "\n".join(f"#{t} — {c} ta" for t, c in top)


def db_delete_note(user_id: int, note_id: int) -> str:
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM notes WHERE id = %s AND user_id = %s", (note_id, user_id))
            deleted = cur.rowcount
    return f"№{note_id} qayd o'chirildi." if deleted else f"№{note_id} qayd topilmadi."


# --- Admin funksiyalar ---

def db_track_user(user_id: int, username: str | None, full_name: str):
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                INSERT INTO users (user_id, username, full_name)
                VALUES (%s, %s, %s)
                ON CONFLICT (user_id) DO UPDATE
                SET username = EXCLUDED.username,
                    full_name = EXCLUDED.full_name,
                    last_seen = NOW(),
                    message_count = users.message_count + 1
            """, (user_id, username, full_name))


def db_get_reply_mode(user_id: int) -> str:
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT reply_mode FROM users WHERE user_id = %s", (user_id,))
            row = cur.fetchone()
    return (row["reply_mode"] if row and row.get("reply_mode") else "text")


def db_set_reply_mode(user_id: int, mode: str):
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                INSERT INTO users (user_id, reply_mode) VALUES (%s, %s)
                ON CONFLICT (user_id) DO UPDATE SET reply_mode = EXCLUDED.reply_mode
            """, (user_id, mode))


def db_get_business_profile(user_id: int) -> dict:
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT business_auto, business_info, business_hours, business_notify "
                "FROM users WHERE user_id = %s",
                (user_id,),
            )
            row = cur.fetchone()
    if not row:
        return {"business_auto": False, "business_info": "", "business_hours": "", "business_notify": True}
    return {
        "business_auto": bool(row.get("business_auto")),
        "business_info": row.get("business_info") or "",
        "business_hours": row.get("business_hours") or "",
        "business_notify": bool(row.get("business_notify", True)),
    }


_BUSINESS_FIELDS = ("business_auto", "business_info", "business_hours", "business_notify")


def db_set_business_field(user_id: int, field: str, value):
    if field not in _BUSINESS_FIELDS:
        raise ValueError(f"Noto'g'ri maydon: {field}")
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(f"""
                INSERT INTO users (user_id, {field}) VALUES (%s, %s)
                ON CONFLICT (user_id) DO UPDATE SET {field} = EXCLUDED.{field}
            """, (user_id, value))


def db_upsert_business_chat(owner_id: int, chat_id: int, conn_id: str, name: str, username: str | None):
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                INSERT INTO business_chats (owner_id, chat_id, conn_id, name, username, last_seen)
                VALUES (%s, %s, %s, %s, %s, NOW())
                ON CONFLICT (owner_id, chat_id) DO UPDATE
                SET conn_id = EXCLUDED.conn_id, name = EXCLUDED.name,
                    username = EXCLUDED.username, last_seen = NOW()
            """, (owner_id, chat_id, conn_id, name, username))


def db_find_business_chats(owner_id: int, query: str) -> list[dict]:
    """Ism yoki username bo'yicha suhbatdoshni qidiradi (oxirgi yozishganlar ichidan)."""
    q = f"%{query.strip().lstrip('@')}%"
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                SELECT chat_id, conn_id, name, username FROM business_chats
                WHERE owner_id = %s AND (name ILIKE %s OR username ILIKE %s)
                ORDER BY last_seen DESC LIMIT 5
            """, (owner_id, q, q))
            return list(cur.fetchall())


def db_list_business_chats(owner_id: int, limit: int = 10) -> list[dict]:
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                SELECT chat_id, conn_id, name, username FROM business_chats
                WHERE owner_id = %s ORDER BY last_seen DESC LIMIT %s
            """, (owner_id, limit))
            return list(cur.fetchall())


def db_business_log_add(owner_id: int, chat_id: int, sender: str, text: str):
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO business_log (owner_id, chat_id, sender, text) VALUES (%s, %s, %s, %s)",
                (owner_id, chat_id, sender, text[:300]),
            )


def db_business_today(owner_id: int) -> list[dict]:
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                SELECT sender, text FROM business_log
                WHERE owner_id = %s AND created_at >= CURRENT_DATE ORDER BY id
            """, (owner_id,))
            return list(cur.fetchall())


def db_business_owners_with_auto() -> list[int]:
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT user_id FROM users WHERE business_auto = TRUE")
            return [r["user_id"] for r in cur.fetchall()]


def db_business_log_cleanup():
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM business_log WHERE created_at < NOW() - INTERVAL '30 days'")


def db_admin_stats() -> str:
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT COUNT(*) as cnt FROM users")
            total_users = cur.fetchone()["cnt"]
            cur.execute("SELECT COUNT(*) as cnt FROM users WHERE last_seen >= NOW() - INTERVAL '24 hours'")
            active_today = cur.fetchone()["cnt"]
            cur.execute("SELECT COUNT(*) as cnt FROM users WHERE last_seen >= NOW() - INTERVAL '7 days'")
            active_week = cur.fetchone()["cnt"]
            cur.execute("SELECT SUM(message_count) as cnt FROM users")
            total_msgs = cur.fetchone()["cnt"] or 0
            cur.execute("SELECT COUNT(*) as cnt FROM reminders WHERE sent = 0")
            active_reminders = cur.fetchone()["cnt"]
    return (
        f"📊 Bot statistikasi:\n\n"
        f"👥 Jami foydalanuvchilar: {total_users}\n"
        f"🟢 Bugun faol: {active_today}\n"
        f"📅 Hafta faol: {active_week}\n"
        f"💬 Jami xabarlar: {total_msgs}\n"
        f"⏰ Faol eslatmalar: {active_reminders}"
    )


def db_admin_users(limit: int = 10) -> str:
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                SELECT user_id, username, full_name, message_count, last_seen
                FROM users ORDER BY last_seen DESC LIMIT %s
            """, (limit,))
            rows = cur.fetchall()
    if not rows:
        return "Foydalanuvchilar yo'q."
    lines = ["👥 Oxirgi foydalanuvchilar:\n"]
    for r in rows:
        name = f"@{r['username']}" if r['username'] else r['full_name']
        last = str(r['last_seen'])[:16]
        lines.append(f"• {name} — {r['message_count']} xabar ({last})")
    return "\n".join(lines)


def db_get_all_user_ids() -> list[int]:
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT user_id FROM users WHERE approved = TRUE")
            return [r["user_id"] for r in cur.fetchall()]


def db_is_approved(user_id: int) -> bool:
    if user_id == ADMIN_ID:
        return True
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT approved FROM users WHERE user_id = %s", (user_id,))
            row = cur.fetchone()
            return bool(row and row["approved"])


def db_approve_user(user_id: int) -> bool:
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("UPDATE users SET approved = TRUE WHERE user_id = %s", (user_id,))
            return cur.rowcount > 0


def db_revoke_user(user_id: int) -> bool:
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("UPDATE users SET approved = FALSE WHERE user_id = %s", (user_id,))
            return cur.rowcount > 0


def db_pending_users() -> list[dict]:
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                SELECT user_id, username, full_name, first_seen
                FROM users WHERE approved = FALSE ORDER BY first_seen DESC LIMIT 20
            """)
            return cur.fetchall()


# --- Profil ---

def db_get_profile(user_id: int) -> dict | None:
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT * FROM profiles WHERE user_id = %s", (user_id,))
            return cur.fetchone()


def db_save_profile(user_id: int, name: str, profession: str, interests: str, goals: str, language: str = "uz"):
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                INSERT INTO profiles (user_id, name, profession, interests, goals, language, onboarded)
                VALUES (%s, %s, %s, %s, %s, %s, TRUE)
                ON CONFLICT (user_id) DO UPDATE SET
                    name = EXCLUDED.name,
                    profession = EXCLUDED.profession,
                    interests = EXCLUDED.interests,
                    goals = EXCLUDED.goals,
                    language = EXCLUDED.language,
                    onboarded = TRUE
            """, (user_id, name, profession, interests, goals, language))


def db_is_onboarded(user_id: int) -> bool:
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT onboarded FROM profiles WHERE user_id = %s", (user_id,))
            row = cur.fetchone()
            return bool(row and row["onboarded"])


def db_add_memory(user_id: int, summary: str):
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO long_memory (user_id, summary) VALUES (%s, %s)",
                (user_id, summary)
            )
            # Faqat oxirgi 10 ta xotirani saqlaymiz
            cur.execute("""
                DELETE FROM long_memory WHERE user_id = %s AND id NOT IN (
                    SELECT id FROM long_memory WHERE user_id = %s ORDER BY id DESC LIMIT 10
                )
            """, (user_id, user_id))


def db_get_memory(user_id: int) -> str:
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT summary FROM long_memory WHERE user_id = %s ORDER BY id DESC LIMIT 10",
                (user_id,)
            )
            rows = cur.fetchall()
    return "\n".join(r["summary"] for r in reversed(rows)) if rows else ""


def db_clear_memory(user_id: int) -> int:
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM long_memory WHERE user_id = %s", (user_id,))
            return cur.rowcount


# --- Suhbat tarixi (deploy/restartda kontekst yo'qolmasligi uchun) ---
# chat_id: Telegram doim 0; ilovada har suhbat o'z chat_id siga ega (ChatGPT uslubi)

def db_save_history_turn(user_id: int, role: str, text: str, chat_id: int = 0):
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO chat_history (user_id, role, text, chat_id) VALUES (%s, %s, %s, %s)",
                (user_id, role, text[:8000], chat_id),
            )
            # Har suhbat uchun faqat oxirgi MAX_HISTORY*2 qatorni saqlaymiz
            cur.execute(
                """DELETE FROM chat_history WHERE user_id = %s AND chat_id = %s AND id NOT IN (
                       SELECT id FROM chat_history WHERE user_id = %s AND chat_id = %s
                       ORDER BY id DESC LIMIT %s
                   )""",
                (user_id, chat_id, user_id, chat_id, _history_limit(chat_id)),
            )


def db_load_history(user_id: int, chat_id: int = 0) -> list:
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """SELECT role, text FROM chat_history
                   WHERE user_id = %s AND chat_id = %s ORDER BY id DESC LIMIT %s""",
                (user_id, chat_id, _history_limit(chat_id)),
            )
            rows = cur.fetchall()
    return [
        types.Content(role=r["role"], parts=[types.Part.from_text(text=r["text"])])
        for r in reversed(rows)
    ]


def db_clear_history(user_id: int, chat_id: int | None = None):
    """chat_id berilsa — faqat o'sha suhbat, berilmasa foydalanuvchining hamma suhbatlari."""
    with get_conn() as conn:
        with conn.cursor() as cur:
            if chat_id is None:
                cur.execute("DELETE FROM chat_history WHERE user_id = %s", (user_id,))
            else:
                cur.execute(
                    "DELETE FROM chat_history WHERE user_id = %s AND chat_id = %s",
                    (user_id, chat_id),
                )


# ============================================================
# YOUTUBE VIDEO TOPISH (to'g'ridan-to'g'ri ijro etish uchun)
# ============================================================

def resolve_youtube_video(query: str) -> str:
    """Google CSE orqali so'ralgan qo'shiq/video uchun aniq YouTube havolasini topadi."""
    if not (GOOGLE_CSE_KEY and GOOGLE_CSE_ID):
        return ""
    try:
        resp = requests.get(
            "https://www.googleapis.com/customsearch/v1",
            params={
                "key": GOOGLE_CSE_KEY, "cx": GOOGLE_CSE_ID,
                "q": f"{query} youtube",
                "num": 5,
            },
            timeout=10,
        )
        items = resp.json().get("items", [])
        for it in items:
            link = it.get("link", "")
            if "youtube.com/watch" in link or "youtu.be/" in link:
                return link
    except Exception:
        logger.exception("YouTube video qidirishda xato")
    return ""


# ============================================================
# INTERNET QIDIRUV
# ============================================================

def do_web_search(query: str) -> str:
    # 1-urinish: Google Custom Search API
    if GOOGLE_CSE_KEY and GOOGLE_CSE_ID:
        try:
            resp = requests.get(
                "https://www.googleapis.com/customsearch/v1",
                params={
                    "key": GOOGLE_CSE_KEY,
                    "cx": GOOGLE_CSE_ID,
                    "q": query,
                    "num": 6,
                },
                timeout=12,
            )
            data = resp.json()
            items = data.get("items", [])
            if not items:
                return _gemini_search(query)

            snippets = "\n\n".join(
                f"[{i+1}] {it.get('title','')}\n{it.get('snippet','')}\nManba: {it.get('link','')}"
                for i, it in enumerate(items)
            )
            now = now_local().strftime("%Y-%m-%d")
            prompt = (
                f"Bugungi sana: {now}. Quyidagi internet qidiruv natijalari asosida "
                f"'{query}' savoliga aniq, qisqa javob ber (o'zbek tilida).\n\n"
                f"{snippets}\n\n"
                f"MUHIM: Faqat natijalardagi ma'lumotga tayan. Agar natijalarda javob bo'lmasa, "
                f"'Bu haqda aniq ma'lumot topilmadi' deb ayt. O'zingdan to'qib chiqarma."
            )
            resp2 = client.models.generate_content(model=MODEL, contents=prompt)
            answer = (resp2.text or "").strip()
            return answer or snippets
        except Exception:
            logger.exception("Google CSE xato")
            return _gemini_search(query)

    return _gemini_search(query)


def _gemini_search(query: str) -> str:
    """Fallback: Gemini grounding bilan qidirish."""
    try:
        response = client.models.generate_content(
            model=MODEL,
            contents=f"Quyidagi savolga internetdagi eng yangi ma'lumotlar asosida qisqa va aniq javob ber: {query}",
            config=types.GenerateContentConfig(
                tools=[types.Tool(google_search=types.GoogleSearch())],
            ),
        )
        if not response.candidates:
            return "Qidiruv natijasi topilmadi."
        try:
            return (response.text or "").strip() or "Ma'lumot topilmadi."
        except Exception:
            parts = response.candidates[0].content.parts if response.candidates[0].content else []
            return " ".join(p.text for p in parts if p.text).strip() or "Ma'lumot topilmadi."
    except Exception as e:
        logger.exception("Gemini search xato")
        return f"Qidiruvda xatolik: {e}"


# ============================================================
# HAVOLA O'QISH (URL kontentini olish)
# ============================================================

import re as _re

_SOCIAL_DOMAINS = ("instagram.com", "tiktok.com", "t.me", "facebook.com", "fb.com", "twitter.com", "x.com")

def do_fetch_url(url: str) -> str:
    if not url.startswith(("http://", "https://")):
        url = "https://" + url

    # Ijtimoiy tarmoqlar login talab qiladi — qidiruvga yo'naltiramiz
    if any(d in url.lower() for d in _SOCIAL_DOMAINS):
        return (
            "Bu ijtimoiy tarmoq sahifasi (Instagram/TikTok/Telegram va h.k.) — "
            "ular login talab qilgani uchun to'g'ridan-to'g'ri o'qib bo'lmaydi. "
            "Iltimos profil/akkaunt nomini (username) yoki mavzuni ayting, men web_search bilan qidiraman."
        )
    try:
        resp = requests.get(
            url,
            headers={
                "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 "
                              "(KHTML, like Gecko) Chrome/120.0 Safari/537.36",
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
                "Accept-Language": "uz,ru;q=0.9,en;q=0.8",
            },
            timeout=12,
            allow_redirects=True,
        )
        if resp.status_code == 403 or resp.status_code == 401:
            return "Bu sayt avtomatik o'qishni bloklagan (403). Mazmunini o'zingiz qisqacha aytsangiz tahlil qilaman."
        ctype = resp.headers.get("content-type", "")
        if "html" not in ctype and "text" not in ctype:
            return f"Bu havola matn sahifa emas ({ctype}). Tahlil qila olmadim."

        html = resp.text
        html = _re.sub(r"<script[\s\S]*?</script>", " ", html, flags=_re.I)
        html = _re.sub(r"<style[\s\S]*?</style>", " ", html, flags=_re.I)
        title_m = _re.search(r"<title[^>]*>(.*?)</title>", html, flags=_re.I | _re.S)
        title = title_m.group(1).strip() if title_m else ""
        text = _re.sub(r"<[^>]+>", " ", html)
        text = _re.sub(r"\s+", " ", text).strip()
        if not text:
            return "Sahifadan matn topilmadi (ehtimol JavaScript bilan yuklanadi). Mazmunini o'zingiz ayting."
        text = text[:6000]
        return f"SAHIFA: {title}\nURL: {url}\n\nMAZMUN:\n{text}"
    except Exception as e:
        logger.exception("URL o'qishda xato")
        return f"Havolani ocholmadim: {e}. Mazmunini o'zingiz qisqacha aytsangiz tahlil qilaman."


# ============================================================
# HUJJAT O'QISH (Word, Excel, matn)
# ============================================================

def extract_docx(data: bytes) -> str:
    import docx
    doc = docx.Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # jadvallarni ham olamiz
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract_xlsx(data: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    out = []
    for ws in wb.worksheets:
        out.append(f"=== Varaq: {ws.title} ===")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                out.append(" | ".join(cells))
            if len(out) > 500:  # juda katta fayllarni cheklash
                out.append("... (qisqartirildi)")
                break
    return "\n".join(out).strip()


# ============================================================
# RASM YARATISH (Imagen)
# ============================================================

IMAGE_MODEL = os.getenv("IMAGE_MODEL", "gemini-3.1-flash-image")  # Nano Banana: yaratish + tahrirlash

# Foydalanuvchining oxirgi yuklagan rasmi — "buni tahrirla" deganda ishlatiladi (15 daqiqa)
last_user_image: dict[int, tuple] = {}  # uid -> (bytes, mime, timestamp)


def _is_quota_error(e: Exception) -> bool:
    m = str(e).lower()
    return "429" in m or "quota" in m or "resource_exhausted" in m or "exceeded" in m


def do_generate_image(prompt: str) -> bytes | None:
    """Matndan rasm yaratadi. Avval Gemini (Nano Banana), bo'lmasa Imagen."""
    try:
        resp = client.models.generate_content(
            model=IMAGE_MODEL,
            contents=[types.Part.from_text(text=prompt)],
            config=types.GenerateContentConfig(response_modalities=["IMAGE", "TEXT"]),
        )
        cand = resp.candidates[0] if resp.candidates else None
        parts = (cand.content.parts or []) if cand and cand.content else []
        for part in parts:
            if part.inline_data and part.inline_data.data:
                return part.inline_data.data
        logger.warning("Rasm yaratish: bo'sh (model=%s)", IMAGE_MODEL)
    except Exception:
        logger.exception("Rasm yaratishda (Gemini) xato — Imagen'ga o'tamiz")
    # Zaxira: Imagen
    try:
        result = client.models.generate_images(
            model="imagen-3.0-generate-002",
            prompt=prompt,
            config=types.GenerateImagesConfig(number_of_images=1),
        )
        if result.generated_images:
            return result.generated_images[0].image.image_bytes
    except Exception:
        logger.exception("Rasm yaratishda (Imagen) xato")
    return None


def do_edit_image(image_bytes: bytes, prompt: str, mime: str = "image/jpeg") -> tuple:
    """Berilgan rasmni ko'rsatma bo'yicha tahrirlaydi. (rasm_bytes|None, izoh_matni) qaytaradi."""
    full_prompt = (
        "You are a precise photo editor. Apply ONLY the change described below to the "
        "provided image. This is a targeted local edit, NOT a regeneration.\n"
        "STRICT RULES:\n"
        "- Change ONLY what is explicitly requested. Do not touch anything else.\n"
        "- Keep the SAME people and faces (identity, features), same pose, same expression, "
        "same clothing (unless the change is about them), same background, same colors, "
        "same lighting, same camera angle, same composition and framing.\n"
        "- Match the original resolution, style and photorealism. The result must look like "
        "the same photo with only the requested edit, not a new image.\n"
        "- Do not add text, watermarks, borders or extra objects.\n"
        "REQUESTED CHANGE: " + prompt
    )
    for attempt in range(2):
        try:
            resp = client.models.generate_content(
                model=IMAGE_MODEL,
                contents=[
                    types.Part.from_bytes(data=image_bytes, mime_type=mime),
                    types.Part.from_text(text=full_prompt),
                ],
                config=types.GenerateContentConfig(response_modalities=["IMAGE", "TEXT"]),
            )
            cand = resp.candidates[0] if resp.candidates else None
            parts = (cand.content.parts or []) if cand and cand.content else []
            out_img, out_text = None, ""
            for part in parts:
                if part.inline_data and part.inline_data.data:
                    out_img = part.inline_data.data
                elif part.text:
                    out_text += part.text
            if out_img:
                return out_img, out_text.strip(), ""
            logger.warning("Rasm tahrir: bo'sh javob (urinish %d)", attempt + 1)
        except Exception as e:
            logger.exception("Rasm tahrirlashda xato")
            if _is_quota_error(e):
                return None, "", "quota"
        time.sleep(1)
    return None, "", ""


def _is_edit_instruction(text: str) -> bool:
    """Matn rasmni TAHRIRLASH buyrug'imi? (savol/tahlil emas)"""
    t = text.lower().translate(_CYR2LAT)
    edit_kw = (
        "o'zgartir", "ozgartir", "tahrir", "tahrirla", "edit", "qo'sh", "qosh",
        "olib tashla", "o'chir", "ochir", "remove", "add", "fon", "background",
        "orqa fon", "rang", "rangini", "color", "uslub", "style", "stil",
        "chiroyli", "chiroyliroq", "yaxshila", "enhance", "sifat", "yorug'",
        "qorong'i", "kattalashtir", "kichiklashtir", "kes", "crop", "ko'zoynak",
        "soch", "kiyim", "ko'ylak", "kulgili", "anime", "rasm qilib", "surat qil",
        "qora oq", "qora-oq", "eskirtir", "yoshartir", "qilib ber", "qilib yubor",
        "ko'k", "kok", "qizil", "yashil", "sariq", "oq qil", "qora qil", "kul rang",
        "chap", "o'ng", "yuqori", "past", "kattaroq", "kichikroq", "yorqin", "xira",
    )
    return any(k in t for k in edit_kw)


def _is_analysis_question(text: str) -> bool:
    """Matn rasmni TUSHUNTIRISH/O'QISH so'rovimi?"""
    t = text.lower().translate(_CYR2LAT)
    q_kw = (
        "nima", "kim", "necha", "qancha", "o'qi", "oqi", "tahlil", "chek",
        "matn", "yozilgan", "tarjima", "nima deb", "ayt", "tushuntir", "bu qanaqa",
    )
    return t.strip().endswith("?") or any(k in t for k in q_kw)


def _mentions_image(text: str) -> bool:
    t = text.lower().translate(_CYR2LAT)
    kw = ("rasm", "surat", "foto", "photo", "image", "logo", "dizayn", "banner",
          "shunday qil", "shunga o'xshash", "shunga oxshash", "buni", "bunga")
    return any(k in t for k in kw)


# ============================================================
# OVOZ YARATISH (Gemini TTS — tabiiy ovoz, o'zbek)
# ============================================================

import struct

TTS_MODEL = os.getenv("TTS_MODEL", "gemini-2.5-flash-preview-tts")
TTS_VOICE = os.getenv("TTS_VOICE", "Kore")  # tabiiy ayol ovozi


def _pcm_to_wav(pcm: bytes, rate: int = 24000, channels: int = 1, bits: int = 16) -> bytes:
    byte_rate = rate * channels * bits // 8
    block_align = channels * bits // 8
    data_size = len(pcm)
    header = (
        b"RIFF" + struct.pack("<I", 36 + data_size) + b"WAVE" +
        b"fmt " + struct.pack("<IHHIIHH", 16, 1, channels, rate, byte_rate, block_align, bits) +
        b"data" + struct.pack("<I", data_size)
    )
    return header + pcm


def do_tts(text: str) -> bytes | None:
    """Matnni tabiiy ovozga aylantiradi (WAV bytes). Bo'sh javobda qayta urinadi."""
    for attempt in range(2):
        try:
            # TTS modelga aniq ko'rsatma beramiz — matnga javob bermasin, faqat o'qisin
            resp = client.models.generate_content(
                model=TTS_MODEL,
                contents=f"Read aloud the following text exactly as written, in a natural warm tone: {text}",
                config=types.GenerateContentConfig(
                    response_modalities=["AUDIO"],
                    speech_config=types.SpeechConfig(
                        voice_config=types.VoiceConfig(
                            prebuilt_voice_config=types.PrebuiltVoiceConfig(voice_name=TTS_VOICE)
                        )
                    ),
                ),
            )
            cand = resp.candidates[0] if resp.candidates else None
            parts = (cand.content.parts or []) if cand and cand.content else []
            for part in parts:
                if part.inline_data and part.inline_data.data:
                    return _pcm_to_wav(part.inline_data.data)
            logger.warning(
                "TTS bo'sh javob (urinish %d). finish_reason=%s",
                attempt + 1, getattr(cand, "finish_reason", None),
            )
        except Exception:
            logger.exception("TTS xato")
        time.sleep(1)
    return None


def _wav_to_ogg(wav: bytes) -> bytes | None:
    """WAV ni Telegram voice bubble uchun OGG/Opus ga o'giradi (ffmpeg bo'lsa)."""
    import subprocess
    try:
        p = subprocess.run(
            ["ffmpeg", "-i", "pipe:0", "-c:a", "libopus", "-b:a", "48k", "-f", "ogg", "pipe:1"],
            input=wav, capture_output=True, timeout=60,
        )
        if p.returncode == 0 and p.stdout:
            return p.stdout
    except Exception:
        pass  # ffmpeg yo'q bo'lsa — WAV'ni audio fayl sifatida yuboramiz
    return None


# ============================================================
# OB-HAVO (Open-Meteo — bepul, kalitsiz, aniq)
# ============================================================

_WEATHER_CODES = {
    0: "ochiq, quyoshli", 1: "asosan ochiq", 2: "qisman bulutli", 3: "bulutli",
    45: "tumanli", 48: "qirovli tuman", 51: "yengil shivalama", 53: "shivalama",
    55: "kuchli shivalama", 56: "muzli shivalama", 57: "kuchli muzli shivalama",
    61: "yengil yomg'ir", 63: "yomg'ir", 65: "kuchli yomg'ir",
    66: "muzli yomg'ir", 67: "kuchli muzli yomg'ir",
    71: "yengil qor", 73: "qor", 75: "kuchli qor", 77: "qor donalari",
    80: "yengil jala", 81: "jala", 82: "kuchli jala",
    85: "yengil qor jalasi", 86: "kuchli qor jalasi",
    95: "momaqaldiroq", 96: "do'lli momaqaldiroq", 99: "kuchli do'lli momaqaldiroq",
}


def do_get_weather(location: str, when: str = "bugun") -> str:
    try:
        geo = requests.get(
            "https://geocoding-api.open-meteo.com/v1/search",
            params={"name": location, "count": 1, "language": "ru"},
            timeout=10,
        ).json()
        results = geo.get("results")
        if not results:
            return f"'{location}' joyi topilmadi. Shahar nomini aniqroq yozing."
        place = results[0]
        lat, lon = place["latitude"], place["longitude"]
        name = place.get("name", location)
        country = place.get("country", "")

        wx = requests.get(
            "https://api.open-meteo.com/v1/forecast",
            params={
                "latitude": lat, "longitude": lon,
                "current": "temperature_2m,relative_humidity_2m,weather_code,wind_speed_10m,apparent_temperature",
                "daily": "temperature_2m_max,temperature_2m_min,weather_code",
                "timezone": "auto", "forecast_days": 3,
            },
            timeout=10,
        ).json()

        cur = wx.get("current", {})
        daily = wx.get("daily", {})
        code = cur.get("weather_code", 0)
        desc = _WEATHER_CODES.get(code, "noma'lum")

        loc_label = f"{name}" + (f", {country}" if country else "")

        if when in ("ertaga", "tomorrow"):
            idx = 1
            day_label = "Ertaga"
        elif when in ("indinga", "after_tomorrow"):
            idx = 2
            day_label = "Indinga"
        else:
            # bugun — joriy ob-havo
            tmax = daily.get("temperature_2m_max", [None])[0]
            tmin = daily.get("temperature_2m_min", [None])[0]
            return (
                f"🌤 {loc_label} — bugun:\n"
                f"Hozir: {cur.get('temperature_2m')}°C (his: {cur.get('apparent_temperature')}°C), {desc}\n"
                f"Kun davomida: {tmin}…{tmax}°C\n"
                f"Namlik: {cur.get('relative_humidity_2m')}%, shamol: {cur.get('wind_speed_10m')} km/soat"
            )

        codes = daily.get("weather_code", [])
        tmax = daily.get("temperature_2m_max", [])
        tmin = daily.get("temperature_2m_min", [])
        if len(codes) > idx:
            d_desc = _WEATHER_CODES.get(codes[idx], "noma'lum")
            return (
                f"🌤 {loc_label} — {day_label.lower()}:\n"
                f"Harorat: {tmin[idx]}…{tmax[idx]}°C, {d_desc}"
            )
        return f"{loc_label} uchun {day_label.lower()} prognozi topilmadi."
    except Exception as e:
        logger.exception("Ob-havo xato")
        return f"Ob-havo ma'lumotini olishda xatolik: {e}"


# ============================================================
# KRIPTO NARXI (CoinGecko — bepul, kalitsiz, real-time)
# ============================================================

def do_get_crypto(coin: str) -> str:
    try:
        # Nomdan coin id ni aniqlash
        search = requests.get(
            "https://api.coingecko.com/api/v3/search",
            params={"query": coin}, timeout=10,
        ).json()
        coins = search.get("coins", [])
        if not coins:
            return f"'{coin}' kriptovalyutasi topilmadi."
        coin_id = coins[0]["id"]
        symbol = coins[0]["symbol"].upper()
        name = coins[0]["name"]

        price = requests.get(
            "https://api.coingecko.com/api/v3/simple/price",
            params={
                "ids": coin_id,
                "vs_currencies": "usd",
                "include_24hr_change": "true",
            },
            timeout=10,
        ).json()
        d = price.get(coin_id, {})
        usd = d.get("usd")
        change = d.get("usd_24h_change", 0)
        if usd is None:
            return f"{name} narxi topilmadi."
        arrow = "📈" if change >= 0 else "📉"
        return (
            f"💰 {name} ({symbol}):\n"
            f"Narx: ${usd:,.2f}\n"
            f"24 soat: {arrow} {change:+.2f}%"
        )
    except Exception as e:
        logger.exception("Kripto narx xato")
        return f"Kripto narxini olishda xatolik: {e}"


# ============================================================
# GEMINI FUNCTION CALLING
# ============================================================

FUNCTION_DECLARATIONS = [
    types.FunctionDeclaration(
        name="web_search",
        description="Internetdan yangi ma'lumot qidirish (yangiliklar, narxlar, valyuta kursi, mashhur odamlar, faktlar). DIQQAT: ob-havo uchun bu emas, get_weather ishlat!",
        parameters=types.Schema(
            type=types.Type.OBJECT,
            properties={"query": types.Schema(type=types.Type.STRING, description="Qidiruv so'rovi. Aniq bo'lsin: agar O'zbekistonga oid bo'lsa 'O'zbekiston' so'zini qo'sh.")},
            required=["query"],
        ),
    ),
    types.FunctionDeclaration(
        name="get_weather",
        description="Ob-havo ma'lumotini olish. 'Ob-havo qanaqa', 'bugun/ertaga havo' kabi savollarda ishlatiladi. Real aniq ma'lumot beradi.",
        parameters=types.Schema(
            type=types.Type.OBJECT,
            properties={
                "location": types.Schema(type=types.Type.STRING, description="Shahar yoki tuman nomi (masalan: Buxoro, G'ijduvon, Toshkent)"),
                "when": types.Schema(type=types.Type.STRING, enum=["bugun", "ertaga", "indinga"], description="Qaysi kun"),
            },
            required=["location"],
        ),
    ),
    types.FunctionDeclaration(
        name="get_crypto",
        description="Kriptovalyuta narxini real-time olish (Bitcoin, Ethereum, Toncoin va boshqalar). 'Bitcoin narxi qancha' kabi savollarda web_search EMAS, SHUNI ishlat.",
        parameters=types.Schema(
            type=types.Type.OBJECT,
            properties={
                "coin": types.Schema(type=types.Type.STRING, description="Kripto nomi (masalan: bitcoin, ethereum, ton, solana)"),
            },
            required=["coin"],
        ),
    ),
    types.FunctionDeclaration(
        name="fetch_url",
        description="Havola (URL/link) mazmunini o'qish va tahlil qilish. Foydalanuvchi link yuborsa yoki 'shu saytni ko'r', 'bu maqolani o'qib ber' desa ishlatiladi.",
        parameters=types.Schema(
            type=types.Type.OBJECT,
            properties={
                "url": types.Schema(type=types.Type.STRING, description="To'liq havola (masalan: https://example.com/page)"),
            },
            required=["url"],
        ),
    ),
    types.FunctionDeclaration(
        name="send_business_message",
        description=(
            "Egasining Telegram suhbatdoshiga (mijoz/kontakt) UNING NOMIDAN haqiqiy xabar yuborish. "
            "Foydalanuvchi 'X ga yozib yubor', 'unga ... deb javob ber', 'xat yubor' desa SHU funksiyani chaqir. "
            "Faqat oldin yozishgan suhbatdoshlarga yuborish mumkin (Telegram Business orqali). "
            "MUHIM: bu funksiyani chaqirmasdan turib 'yubordim' dema!"
        ),
        parameters=types.Schema(
            type=types.Type.OBJECT,
            properties={
                "recipient": types.Schema(
                    type=types.Type.STRING,
                    description="Suhbatdoshning ismi yoki @username (suhbat kontekstidan yoki foydalanuvchi aytganidan ol)",
                ),
                "message": types.Schema(type=types.Type.STRING, description="Yuboriladigan xabar matni"),
            },
            required=["recipient", "message"],
        ),
    ),
    types.FunctionDeclaration(
        name="remember",
        description="Foydalanuvchi haqida MUHIM, uzoq muddat eslab qolish kerak bo'lgan faktni saqlash. Masalan: uning loyihasi, ishi, maqsadi, sevimli narsasi, oilasi, muhim sanalari, qarorlari. Foydalanuvchi o'zi haqida yangi muhim narsa aytsa — DARHOL shuni chaqir.",
        parameters=types.Schema(
            type=types.Type.OBJECT,
            properties={
                "fact": types.Schema(type=types.Type.STRING, description="Eslab qolinadigan qisqa fakt (masalan: 'Foydalanuvchi YouTube blog ochmoqchi, mavzu - sayohat')"),
            },
            required=["fact"],
        ),
    ),
    types.FunctionDeclaration(
        name="generate_image",
        description="Rasm/surat yaratish. 'Rasm chiz', 'surat yaratib ber', 'menga ... rasmini chiz' desa ishlatiladi. Prompt ingliz tilida va batafsil bo'lsa sifat yuqori bo'ladi.",
        parameters=types.Schema(
            type=types.Type.OBJECT,
            properties={
                "prompt": types.Schema(type=types.Type.STRING, description="Rasm tavsifi (ingliz tilida, batafsil: uslub, rang, kompozitsiya)"),
            },
            required=["prompt"],
        ),
    ),
    types.FunctionDeclaration(
        name="add_transaction",
        description="Kirim yoki chiqimni bazaga yozish. Pul sarflagani/topgani haqida aytsa yoki chek rasmida summa ko'rinsa ishlatiladi.",
        parameters=types.Schema(
            type=types.Type.OBJECT,
            properties={
                "tx_type": types.Schema(type=types.Type.STRING, enum=["kirim", "chiqim"]),
                "amount": types.Schema(type=types.Type.NUMBER, description="Summa so'mda. '50 ming' = 50000"),
                "category": types.Schema(type=types.Type.STRING, description="oziq-ovqat, transport, kommunal, maosh, savdo, boshqa..."),
                "note": types.Schema(type=types.Type.STRING, description="Qisqa izoh (ixtiyoriy)"),
            },
            required=["tx_type", "amount", "category"],
        ),
    ),
    types.FunctionDeclaration(
        name="get_report",
        description="Kirim-chiqim hisoboti. 'Qancha sarfladim', 'hisobot', 'balans' desa ishlatiladi. Aniq sana oralig'i so'ralsa (masalan '1-iyundan 10-iyungacha', 'shu oyning 5-sanasi') start_date/end_date ber.",
        parameters=types.Schema(
            type=types.Type.OBJECT,
            properties={
                "period": types.Schema(type=types.Type.STRING, enum=["bugun", "kecha", "hafta", "oy"], description="Tayyor davr. Aniq sana berilsa bo'sh qoldir."),
                "start_date": types.Schema(type=types.Type.STRING, description="Boshlanish sanasi 'YYYY-MM-DD' (ixtiyoriy)"),
                "end_date": types.Schema(type=types.Type.STRING, description="Tugash sanasi 'YYYY-MM-DD' (ixtiyoriy, berilmasa start_date kuni)"),
            },
        ),
    ),
    types.FunctionDeclaration(
        name="delete_last_transaction",
        description="Oxirgi kirim/chiqim yozuvini o'chirish ('xato yozdim', 'oxirgisini o'chir').",
        parameters=types.Schema(type=types.Type.OBJECT, properties={}),
    ),
    types.FunctionDeclaration(
        name="set_reminder",
        description="Eslatma o'rnatish. 'Ertaga 9 da ... eslatib qo'y' kabi so'rovlarda. Hozirgi sana-vaqt system promptda berilgan — 'ertaga', 'bir soatdan keyin' kabilarni o'zing aniq vaqtga aylantir.",
        parameters=types.Schema(
            type=types.Type.OBJECT,
            properties={
                "text": types.Schema(type=types.Type.STRING, description="Nimani eslatish kerak"),
                "remind_at": types.Schema(type=types.Type.STRING, description="Vaqt 'YYYY-MM-DD HH:MM' formatida"),
            },
            required=["text", "remind_at"],
        ),
    ),
    types.FunctionDeclaration(
        name="list_reminders",
        description="Faol eslatmalar ro'yxatini ko'rsatish.",
        parameters=types.Schema(type=types.Type.OBJECT, properties={}),
    ),
    types.FunctionDeclaration(
        name="delete_reminder",
        description="Eslatmani raqami (№) bo'yicha o'chirish/bekor qilish.",
        parameters=types.Schema(
            type=types.Type.OBJECT,
            properties={"reminder_id": types.Schema(type=types.Type.INTEGER, description="Eslatma raqami")},
            required=["reminder_id"],
        ),
    ),
    types.FunctionDeclaration(
        name="add_note",
        description=(
            "Qayd yoki KOD SNIPPETI saqlash. 'Eslab qol: ...', 'saqla', 'snippet qilib qo'y' desa ishlatiladi "
            "(sozlamalar, buyruqlar, kod parchalari, parollar, raqamlar, fikrlar). "
            "Xabarda kod bo'lsa — uni 'code' ga ajratib ber, tushuntirishni 'text' ga. "
            "Teglarni foydalanuvchi #teg deb yozgan bo'lsa yoki mavzudan ravshan bo'lsa (masalan railway, python, kotlin) qo'sh."
        ),
        parameters=types.Schema(
            type=types.Type.OBJECT,
            properties={
                "text": types.Schema(type=types.Type.STRING, description="Tushuntirish/matn (kod bo'lmagan qism)"),
                "title": types.Schema(type=types.Type.STRING, description="Qisqa sarlavha (3-6 so'z), keyin topish oson bo'lishi uchun"),
                "tags": types.Schema(type=types.Type.STRING, description="Vergul bilan teglar, masalan: 'railway,ffmpeg' (# belgisiz)"),
                "code": types.Schema(type=types.Type.STRING, description="Kod/buyruq/konfig parchasi — aynan, o'zgartirmasdan"),
                "lang": types.Schema(type=types.Type.STRING, description="Kod tili: python, kotlin, bash, toml, sql, js... (ixtiyoriy)"),
            },
        ),
    ),
    types.FunctionDeclaration(
        name="find_notes",
        description=(
            "Saqlangan qaydlar/snippetlardan qidirish. 'Railway ffmpeg qanday edi?', 'kotlin snippetlarim', "
            "'#python qaydlar' kabi so'rovlarda. Matn, sarlavha, kod va teglar bo'yicha qidiradi. Bo'sh = oxirgilari."
        ),
        parameters=types.Schema(
            type=types.Type.OBJECT,
            properties={
                "query": types.Schema(type=types.Type.STRING, description="Qidiruv so'zi (masalan: ffmpeg)"),
                "tag": types.Schema(type=types.Type.STRING, description="Faqat shu teg bo'yicha (masalan: railway)"),
            },
        ),
    ),
    types.FunctionDeclaration(
        name="get_note",
        description="Bitta qaydni raqami (№) bo'yicha TO'LIQ ko'rsatish — kod bloki bilan. '№12 ni ko'rsat', '12-qaydni ochib ber'.",
        parameters=types.Schema(
            type=types.Type.OBJECT,
            properties={"note_id": types.Schema(type=types.Type.INTEGER)},
            required=["note_id"],
        ),
    ),
    types.FunctionDeclaration(
        name="list_tags",
        description="Foydalanuvchining barcha teglari va har birida nechta qayd borligi. 'Teglarim', 'qanday kategoriyalar bor' desa.",
        parameters=types.Schema(type=types.Type.OBJECT, properties={}),
    ),
    types.FunctionDeclaration(
        name="delete_note",
        description="Qaydni raqami (№) bo'yicha o'chirish.",
        parameters=types.Schema(
            type=types.Type.OBJECT,
            properties={"note_id": types.Schema(type=types.Type.INTEGER)},
            required=["note_id"],
        ),
    ),
]

# Biznes suhbatlar (begona suhbatdoshlar) uchun faqat XAVFSIZ, o'qish-only funksiyalar:
# qidiruv, ob-havo, kurs. Egasining moliya/eslatma/qaydlariga tegadiganlar KIRMAYDI.
_SAFE_BUSINESS_TOOL_NAMES = {"web_search", "get_weather", "get_crypto", "fetch_url"}
SAFE_BUSINESS_DECLARATIONS = [d for d in FUNCTION_DECLARATIONS if d.name in _SAFE_BUSINESS_TOOL_NAMES]

# Faqat Shoxa (Android ilova) orqali ishlaganda yoqiladi — Telegram'da emas,
# chunki bular HAQIQIY telefonni boshqaradi, buni faqat ilova bajara oladi.
DEVICE_ACTION_DECLARATIONS = [
    types.FunctionDeclaration(
        name="open_app",
        description="Telefondagi ilovani ochish. 'Telegram och', 'YouTube kir', 'Instagram ochib ber' kabi so'rovlarda ishlatiladi.",
        parameters=types.Schema(
            type=types.Type.OBJECT,
            properties={
                "app_name": types.Schema(
                    type=types.Type.STRING,
                    description="Ilova nomi: telegram, youtube, instagram, whatsapp, chrome, gmail, maps, camera, settings, spotify, tiktok, facebook va h.k.",
                ),
            },
            required=["app_name"],
        ),
    ),
    types.FunctionDeclaration(
        name="set_alarm",
        description="Telefonda budilnik/alarm o'rnatish. 'Ertalab 7 da budilnik qo'y', 'juma kuni 6 da uyg'otib qo'y' kabi so'rovlarda ishlatiladi.",
        parameters=types.Schema(
            type=types.Type.OBJECT,
            properties={
                "hour": types.Schema(type=types.Type.INTEGER, description="Soat, 0-23 formatida"),
                "minute": types.Schema(type=types.Type.INTEGER, description="Minut, 0-59"),
                "label": types.Schema(type=types.Type.STRING, description="Budilnik nomi (ixtiyoriy)"),
                "date": types.Schema(
                    type=types.Type.STRING,
                    description="Qaysi kunga — YYYY-MM-DD formatda (ixtiyoriy). 'Ertaga', 'juma kuni' kabi nisbiy kunlarni HOZIRGI VAQT asosida aniq sanaga aylantirib ber. Bugungi/eng yaqin vaqt bo'lsa bo'sh qoldir.",
                ),
            },
            required=["hour", "minute"],
        ),
    ),
    types.FunctionDeclaration(
        name="make_call",
        description="Qo'ng'iroq qilish. 'Onamga qo'ng'iroq qil', '+998901234567 ga qo'ng'iroq qil' kabi so'rovlarda ishlatiladi. Agar ism aytilsa (raqam emas) — contact_name ber, ilova telefondagi kontaktlardan o'zi qidiradi.",
        parameters=types.Schema(
            type=types.Type.OBJECT,
            properties={
                "phone_number": types.Schema(type=types.Type.STRING, description="To'liq telefon raqami, agar aytilgan bo'lsa"),
                "contact_name": types.Schema(type=types.Type.STRING, description="Kontakt ismi, agar raqam o'rniga ism aytilgan bo'lsa (masalan: Ona, Ali)"),
            },
        ),
    ),
    types.FunctionDeclaration(
        name="search_in_app",
        description="Ilova ichida biror narsani qidirish/ochish. 'YouTube'da bu qo'shiqni qidir', 'Yandex Mapsda restoran top' kabi 'ilova och VA shu ishni qil' birikma so'rovlarida ishlatiladi. FAQAT YANGI qidiruv so'ralganda chaqir — agar foydalanuvchi 'eshitamiz', 'xa', 'rahmat', 'zo'r' kabi oddiy javob/tasdiq aytsa, BU FUNKSIYANI QAYTA CHAQIRMA, shunchaki tabiiy javob ber.",
        parameters=types.Schema(
            type=types.Type.OBJECT,
            properties={
                "app_name": types.Schema(type=types.Type.STRING, description="Ilova nomi: youtube, instagram, chrome, maps va h.k."),
                "query": types.Schema(type=types.Type.STRING, description="Qidiriladigan narsa (qo'shiq nomi, joy, mavzu)"),
            },
            required=["app_name", "query"],
        ),
    ),
    types.FunctionDeclaration(
        name="open_telegram_chat",
        description=(
            "Telegram'da ma'lum bir kishi, kanal yoki guruhni ochish, ixtiyoriy ravishda xabar matnini tayyorlab qo'yish. "
            "Quyidagilarning HAMMASIDA ishlatiladi: 'Telegram'dan <kanal>ga kir', '<kishi> bilan suhbatni och', "
            "'Telegram'da <kimga> yoz: ...', '<kanal>ni och'. "
            "Agar @username aniq aytilmasa, kanal/kishi nomini username sifatida yoz (masalan 'Mavzu' kanali -> 'mavzu'). "
            "Foydalanuvchi baribir Yuborish tugmasini bosishi kerak (xabar avtomatik yuborilmaydi) — buni tabiiy ayt."
        ),
        parameters=types.Schema(
            type=types.Type.OBJECT,
            properties={
                "username": types.Schema(type=types.Type.STRING, description="Telegram username (@ belgisiz) yoki kontakt ismi"),
                "message": types.Schema(type=types.Type.STRING, description="Tayyorlanadigan xabar matni (ixtiyoriy)"),
            },
            required=["username"],
        ),
    ),
    types.FunctionDeclaration(
        name="send_sms",
        description="SMS xabar yuborish. 'Bu raqamga SMS yubor: ...' kabi so'rovlarda ishlatiladi. Foydalanuvchi yuborishni xabar ilovasida tasdiqlashi kerak bo'ladi.",
        parameters=types.Schema(
            type=types.Type.OBJECT,
            properties={
                "phone_number": types.Schema(type=types.Type.STRING, description="Qabul qiluvchi raqam"),
                "message": types.Schema(type=types.Type.STRING, description="Xabar matni"),
            },
            required=["phone_number", "message"],
        ),
    ),
    types.FunctionDeclaration(
        name="set_volume",
        description="Telefon ovoz balandligini boshqarish. 'Ovozni baland qil', 'tovushni kamaytir', 'ovozni o'chir' kabi so'rovlarda ishlatiladi.",
        parameters=types.Schema(
            type=types.Type.OBJECT,
            properties={
                "direction": types.Schema(type=types.Type.STRING, enum=["up", "down", "mute"], description="Ovoz yo'nalishi"),
            },
            required=["direction"],
        ),
    ),
    types.FunctionDeclaration(
        name="toggle_flashlight",
        description="Telefon fonarini (chiroq) yoqish/o'chirish. 'Fonarni yoq', 'chiroqni o'chir' kabi so'rovlarda ishlatiladi.",
        parameters=types.Schema(
            type=types.Type.OBJECT,
            properties={
                "on": types.Schema(type=types.Type.BOOLEAN, description="True = yoq, False = o'chir"),
            },
            required=["on"],
        ),
    ),
]


def build_system_prompt(user_id: int | None = None) -> str:
    now = now_local()
    weekdays = ["Dushanba", "Seshanba", "Chorshanba", "Payshanba", "Juma", "Shanba", "Yakshanba"]

    profile_section = ""
    memory_section = ""
    if user_id:
        profile = db_get_profile(user_id)
        if profile:
            profile_section = f"""
FOYDALANUVCHI PROFILI:
- Ismi: {profile['name']}
- Kasbi: {profile['profession']}
- Qiziqishlari: {profile['interests']}
- Maqsadlari: {profile['goals']}
- Tili: {profile['language']}
Foydalanuvchiga doim ismi bilan murojaat qil. Kasbiga mos maslahatlar ber."""

        memory = db_get_memory(user_id)
        if memory:
            memory_section = f"\nOLDINGI SUHBATLARDAN XOTIRA:\n{memory}"

    return f"""Sen Buxorolik dasturchi tomonidan yaratilgan shaxsiy AI agentsan. Sirdosh deb ataladigan Telegram botda ishlaysan.
Agar "sen kimsan", "sen nimasan", "kim yaratdi", "qaysi model", "qaysi AI" kabi savollar berilsa — "Men Buxorolik dasturchi tomonidan yaratilgan shaxsiy AI agentman 😊" deb javob ber. Hech qachon Google yoki boshqa kompaniya nomini aytma.
Sen ovozli xabarlarni ham qabul qila olasan va tushunasan. Foydalanuvchi ovoz yuborsa, uni tushunib javob berasan.
Sen OVOZLI JAVOB ham bera olasan — tizim sening matningni avtomatik ovozga aylantirib yuboradi. Shuning uchun "ovozli javob bera olmayman" deb HECH QACHON aytma. Foydalanuvchi ovozli javob so'rasa — oddiy matn javobingni yozaver, qolganini tizim o'zi qiladi (javob turi /sozlamalar orqali ham tanlanadi).

EGASI NOMIDAN JAVOB YOZISH (juda muhim!):
- Foydalanuvchi birovning xabarini ko'rsatib "mening o'rnimga javob yozib ber", "shunga javob yoz", "shu xabarga javob tayyorla" desa — FAQAT TAYYOR JAVOB MATNINI yoz, boshqa HECH NARSA yozma.
- "Mana javob:", "Quyidagicha javob berishingiz mumkin:", izoh, variantlar, tushuntirish — TAQIQLANADI. Foydalanuvchi sening xabaringni to'g'ridan-to'g'ri nusxalab/forward qilib yuboradi, shuning uchun javob yuborishga tayyor holda bo'lsin.
- Javob foydalanuvchi nomidan (birinchi shaxsda) yoziladi, uning uslubi va vaziyatga mos ohangda.
- Ovozda so'ralsa ham xuddi shunday faqat javob matnini yoz — tizim o'zi ovozga aylantiradi.
- Umumiy qoida: barcha javoblaring keraksiz kirish gaplarsiz, to'g'ridan-to'g'ri va aniq bo'lsin.

HOZIRGI VAQT: {now.strftime('%Y-%m-%d %H:%M')}, {weekdays[now.weekday()]} (Asia/Tashkent).
"Ertaga" = {(now + timedelta(days=1)).strftime('%Y-%m-%d')}. Nisbiy vaqtlarni shu asosda hisobla.

Imkoniyatlaring:
1. Internet qidiruv (web_search) — yangi ma'lumot kerak bo'lsa taxmin qilma, qidir! O'zbekistonga oid odam/joy bo'lsa qidiruvga "O'zbekiston" qo'sh.
2. Ob-havo (get_weather) — ob-havo so'ralsa SHU funksiyani ishlat, web_search EMAS. Joy aytilmasa, profildagi joyni yoki "Toshkent" ni ol.
   Kripto narxi (get_crypto) — Bitcoin/Ethereum kabi narxlar so'ralsa SHU funksiyani ishlat, web_search EMAS (real-time aniq narx).
   Rasm yaratish (generate_image) — "rasm chiz", "surat yaratib ber" desa ishlatiladi. Prompt ni ingliz tilida, JUDA batafsil yoz (uslub, yorug'lik, kompozitsiya, kayfiyat, sifat). Rasm TAHRIRLASH (foydalanuvchi rasm yuklab "buni o'zgartir/fon/rang..." desa) — buni tizim avtomatik bajaradi, sen aralashma.
   Hujjat (PDF/Word/Excel) — foydalanuvchi fayl yuborsa avtomatik o'qiysan va tahlil qilasan.
3. Buxgalteriya — xarajat/daromad aytilsa add_transaction. Hisobot so'ralsa get_report.
4. Eslatmalar — set_reminder (vaqtni aniq 'YYYY-MM-DD HH:MM' ga aylantir).
5. Qaydlar va SNIPPET KUTUBXONASI — "eslab qol", "saqla", "snippet qil" desa add_note: xabarda kod/buyruq/konfig bo'lsa uni AYNAN, o'zgartirmasdan code'ga, tushuntirishni text'ga, qisqa title va mavzuga mos teglar (foydalanuvchi #teg yozgan bo'lsa shuni ol) ber. "... qanday edi?", "... snippetim", "#teg qaydlar" desa find_notes (query va/yoki tag). "№N ni ko'rsat" desa get_note. "teglarim" desa list_tags. Qayd natijasini (ayniqsa kod bloklarini) O'ZGARTIRMASDAN, qanday kelgan bo'lsa shundayligicha foydalanuvchiga yetkaz — u nusxalab ishlatadi.
7. Suhbatdoshga xabar yuborish (send_business_message) — foydalanuvchi "X ga yozib yubor", "unga javob yubor", "xat yubor" desa DARHOL SHU funksiyani chaqir (Telegram Business orqali uning nomidan haqiqiy xabar ketadi). Foydalanuvchi aytgan nomni (masalan "developer") recipient sifatida to'g'ridan-to'g'ri ber — o'zingcha "bu kim ekan" deb mulohaza yuritma, funksiya o'zi topadi yoki ro'yxatni qaytaradi. QAT'IY: funksiya chaqirmasdan turib "yubordim" DEMA — bu yolg'on bo'ladi.
6. Rasmlar — chek/kvitansiya rasmi kelsa, summa va do'konni aniqlab add_transaction chaqir.

TELEFONNI BOSHQARISH (agar shu funksiyalar mavjud bo'lsa — Shoxa ilovasidasan):
- "... och", "... kir", "... ochib ber" + ilova nomi → open_app. HECH IKKILANMA, albatta chaqir, faqat gapirib qo'ymagin.
- Ilova ochish VA shu ilova ichida biror narsa qilish birga aytilsa (masalan "YouTube'da Daler Mansurov qo'shig'ini qidir") → search_in_app (app_name + query), open_app EMAS.
- "Telegram'da <kimga> yoz/xabar yubor" → open_telegram_chat (username + message). Agar shaxs ismi aytilsa lekin @username noma'lum bo'lsa, contact_name sifatida ismni username maydoniga yoz.
- "Telegram'dan <kanal/guruh>ga kir", "<kanal>ni och" → open_telegram_chat (kanal nomini username sifatida ber).
- Qo'ng'iroq: raqam aytilsa phone_number, ism aytilsa (masalan "Onamga qo'ng'iroq qil") contact_name bilan make_call.
- Bularning barchasi HAQIQIY telefonda amalga oshadi — sen faqat signal berasan, natijani "amalga oshirilmoqda" deb tabiiy ayt, "men buni qila olmayman" demagin.
- MUHIM: bir amal (open_app/search_in_app/open_telegram_chat/make_call/set_alarm) bajarilgandan SO'NG, foydalanuvchi "xa", "rahmat", "zo'r", "eshitamiz", "yaxshi" kabi oddiy javob/tasdiq aytsa — HECH QANDAY funksiya chaqirma, faqat tabiiy, qisqa javob ber ("Marhamat!", "Yaxshi tinglang!" va h.k.). Faqat foydalanuvchi YANGI, aniq buyruq bersa qayta funksiya chaqir.

Qoidalar:
- Foydalanuvchi qaysi tilda gapirsa, o'sha tilda javob ber (asosan o'zbek).
- Javoblar qisqa va aniq. Aniq bilmasang "aniq ma'lumot topa olmadim" deb ayt — YOLG'ON to'qima!
- Foydalanuvchi avvalgi xabariga "ha", "yo'q" desa — kontekstni esla, qayta so'rama.
- Summalar: "50 ming" = 50000, "1.5 mln" = 1500000.
- Funksiya natijasini chiroyli, tushunarli qilib yetkaz.

MUHIM — SUHBAT SIFATI:
- Sen HAR SOHADA bilimli aqlli maslahatchisan: blogerlik, biznes, dasturlash, ta'lim, sog'liq, psixologiya, marketing, din, tarix, fan — istalgan mavzuda foydali javob ber.
- HECH QACHON bir xil umumiy javobni takrorlama ("men yordam bera olaman" kabi). Har savolga ANIQ, MAZMUNLI, AMALIY javob ber.
- Biror sohada yordam so'ralsa — umumiy gap urma, KONKRET maslahat, qadam-baqadam reja, aniq misollar ber.
- Foydalanuvchi link/havola yuborsa — fetch_url bilan o'qib, mazmunini tahlil qil. Agar fetch_url ishlamasa (ijtimoiy tarmoq/bloklangan), username yoki mavzuni web_search bilan qidirib top.
- Faqat zarur bo'lganda savol ber. Yetarli ma'lumot bo'lsa — darrov foydali javob ber.
- Sen oddiy "yordamchi" emas, haqiqiy aqlli SIRDOSHsan — chuqur, foydali, inson kabi muloqot qil.

XOTIRA VA G'OYA (eng muhim!):
- Foydalanuvchi o'zi haqida muhim narsa aytsa (ishi, loyihasi, maqsadi, qiziqishi, muammosi, qarori) — DARHOL "remember" funksiyasini chaqirib eslab qol. Keyingi suhbatlarda shuni hisobga ol.
- Foydalanuvchining ishini, loyihasini o'rgan va PROAKTIV ravishda foydali G'OYALAR, takliflar ber — so'ramasa ham. Masalan blogger bo'lsa kontent g'oyalari, tadbirkor bo'lsa biznes takliflari.
- Avvalgi suhbatlardagi xotirani eslab, "o'tgan safar aytgan loyihangiz qanday ketyapti?" kabi tabiiy, g'amxo'r muloqot qil.
- Sen egangning eng yaqin sirdoshi, maslahatchisi va ilhomchisisan.

FOYDALANUVCHINI TINGLA (juda muhim!):
- Foydalanuvchi biror narsani "kerakmas", "boshqasini ayt", "bu menga to'g'ri kelmaydi" desa — DARHOL o'sha mavzuni TASHLA va boshqa, YANGI yo'nalishda fikr ber. Eski mavzuni qayta-qayta takrorlama!
- Uning xohish va e'tirozlarini hurmat qil. Agar bir taklif yoqmasa, butunlay boshqacha variant taklif qil.
- O'zingdan bitta mavzuga yopishib olma — foydalanuvchi nima xohlayotganini diqqat bilan tingla va shunga moslash.
{profile_section}{memory_section}
"""


def execute_function(user_id: int, name: str, args: dict) -> str:
    try:
        if name == "web_search":
            return do_web_search(args.get("query", ""))
        if name == "get_weather":
            return do_get_weather(args.get("location", "Toshkent"), args.get("when", "bugun"))
        if name == "get_crypto":
            return do_get_crypto(args.get("coin", "bitcoin"))
        if name == "fetch_url":
            return do_fetch_url(args.get("url", ""))
        if name == "remember":
            fact = args.get("fact", "").strip()
            if fact:
                db_add_memory(user_id, fact)
                return "Eslab qoldim ✅"
            return "Eslab qolinadigan narsa yo'q."
        if name == "add_transaction":
            return db_add_transaction(
                user_id, args.get("tx_type", "chiqim"), float(args.get("amount", 0)),
                args.get("category", "boshqa"), args.get("note", ""),
            )
        if name == "get_report":
            return db_get_report(
                user_id, args.get("period", "oy"),
                args.get("start_date", ""), args.get("end_date", ""),
            )
        if name == "delete_last_transaction":
            return db_delete_last(user_id)
        if name == "set_reminder":
            return db_set_reminder(user_id, args.get("text", "Eslatma"), args.get("remind_at", ""))
        if name == "list_reminders":
            return db_list_reminders(user_id)
        if name == "delete_reminder":
            return db_delete_reminder(user_id, int(args.get("reminder_id", 0)))
        if name == "add_note":
            if not (args.get("text") or args.get("code")):
                return "Xato: saqlash uchun matn yoki kod bo'sh."
            return db_add_note(
                user_id, args.get("text", ""), args.get("title", ""),
                args.get("tags", ""), args.get("code", ""), args.get("lang", ""),
            )
        if name == "find_notes":
            return db_find_notes(user_id, args.get("query", ""), args.get("tag", ""))
        if name == "get_note":
            return db_get_note(user_id, int(args.get("note_id", 0)))
        if name == "list_tags":
            return db_list_tags(user_id)
        if name == "delete_note":
            return db_delete_note(user_id, int(args.get("note_id", 0)))
        return f"Noma'lum funksiya: {name}"
    except Exception as e:
        logger.exception("Funksiyada xato: %s", name)
        return f"Xatolik: {e}"


# ============================================================
# AGENT SIKLI
# ============================================================

MAX_HISTORY = 20
MAX_HISTORY_PROJECT = 60  # loyiha rejimi: uzun strategik suhbat uchun


def _history_limit(chat_id: int) -> int:
    return (MAX_HISTORY_PROJECT if chat_id == 777 else MAX_HISTORY) * 2
# Kalit: (user_id, chat_id). Telegram uchun chat_id doim 0, ilovada har suhbat alohida.
chat_history: dict[tuple[int, int], list[types.Content]] = {}
onboarding_state: dict[int, dict] = {}  # {user_id: {step, name, profession, interests, goals}}

FALLBACK_MODEL = os.getenv("GEMINI_FALLBACK_MODEL", "gemini-3.5-flash")


def gemini_generate(model: str, contents, config=None):
    """Gemini chaqiruvi: 2 urinish asosiy model, oxirgisi zaxira model bilan.
    Xato ham, bo'sh javob ham qayta uriniladi — "Xatolik yuz berdi" kamayadi."""
    last_exc: Exception | None = None
    response = None
    for attempt in range(3):
        use_model = model if attempt < 2 else FALLBACK_MODEL
        try:
            response = client.models.generate_content(
                model=use_model, contents=contents, config=config,
            )
            if response.candidates:
                return response
            logger.warning("Gemini bo'sh javob (%s, urinish %d)", use_model, attempt + 1)
        except Exception as e:
            last_exc = e
            logger.warning("Gemini xato (%s, urinish %d): %s", use_model, attempt + 1, e)
        if attempt < 2:
            time.sleep(1 + attempt)
    if response is not None:
        return response
    raise last_exc


def _save_history(user_id: int, user_parts: list[types.Part], answer: str, chat_id: int = 0):
    """Foydalanuvchi xabari va javobni tarixga saqlaydi (kontekst saqlanishi uchun)."""
    key = (user_id, chat_id)
    history = chat_history.setdefault(key, [])
    saved = [p if p.text else types.Part.from_text(text="[media xabar]") for p in user_parts]
    history.append(types.Content(role="user", parts=saved))
    history.append(types.Content(role="model", parts=[types.Part.from_text(text=answer)]))
    lim = _history_limit(chat_id)
    if len(history) > lim:
        chat_history[key] = history[-lim:]
    # Postgres'ga ham yozamiz — deploy/restartda kontekst yo'qolmasligi uchun
    try:
        user_text = " ".join(p.text for p in saved if p.text).strip() or "[media xabar]"
        db_save_history_turn(user_id, "user", user_text, chat_id)
        db_save_history_turn(user_id, "model", answer, chat_id)
    except Exception:
        logger.exception("Tarixni bazaga yozishda xato")


async def ask_agent(
    user_id: int,
    user_parts: list[types.Part],
    image_sink: list | None = None,
    device_action_sink: list | None = None,
    chat_id: int = 0,
    system_prompt: str | None = None,
    allow_tools: bool = True,
    tools_override: list | None = None,
) -> str:
    # Birinchi murojaatda (yoki restartdan keyin) tarixni bazadan tiklaymiz
    key = (user_id, chat_id)
    if key not in chat_history:
        try:
            chat_history[key] = await asyncio.to_thread(db_load_history, user_id, chat_id)
        except Exception:
            logger.exception("Tarixni bazadan yuklashda xato")
            chat_history[key] = []
    history = chat_history[key]
    contents = history + [types.Content(role="user", parts=user_parts)]

    # device_action_sink berilgan bo'lsa — Shoxa ilovasidan kelgan so'rov,
    # telefonni boshqarish funksiyalarini ham yoqamiz.
    # allow_tools=False — biznes avto-javob kabi begona suhbatdoshlar uchun:
    # ular egasining moliya/eslatma funksiyalariga tega olmasligi kerak.
    if tools_override is not None:
        declarations = tools_override
    else:
        declarations = (
            FUNCTION_DECLARATIONS + (DEVICE_ACTION_DECLARATIONS if device_action_sink is not None else [])
        ) if allow_tools else []

    config = types.GenerateContentConfig(
        system_instruction=system_prompt or build_system_prompt(user_id),
        temperature=0.7,
        max_output_tokens=2048,
        thinking_config=types.ThinkingConfig(thinking_budget=0),
        tools=[types.Tool(function_declarations=declarations)] if declarations else None,
    )

    answer = "Kechirasiz, javob topa olmadim. Boshqacharoq so'rab ko'ring."
    for _ in range(6):
        response = await asyncio.to_thread(
            gemini_generate, model=MODEL, contents=contents, config=config,
        )
        if not response.candidates:
            logger.warning("ask_agent: bo'sh candidates. prompt_feedback=%s", getattr(response, "prompt_feedback", None))
            break
        candidate = response.candidates[0]
        parts = (candidate.content.parts or []) if candidate.content else []
        if not parts:
            try:
                txt = (response.text or "").strip()
            except Exception:
                txt = ""
            if txt:
                answer = txt
            else:
                logger.warning(
                    "ask_agent: bo'sh parts. finish_reason=%s safety=%s",
                    getattr(candidate, "finish_reason", None),
                    getattr(candidate, "safety_ratings", None),
                )
            break

        function_calls = [p.function_call for p in parts if p.function_call]
        if not function_calls:
            try:
                txt = (response.text or "").strip()
            except Exception:
                txt = ""
            if txt:
                answer = txt
            else:
                logger.warning(
                    "ask_agent: matn yo'q. finish_reason=%s safety=%s parts=%s",
                    getattr(candidate, "finish_reason", None),
                    getattr(candidate, "safety_ratings", None),
                    parts,
                )
            break

        contents.append(candidate.content)
        result_parts = []
        for fc in function_calls:
            args = dict(fc.args or {})
            if fc.name == "send_business_message":
                recipient = str(args.get("recipient", "")).strip()
                msg_text = str(args.get("message", "")).strip()
                if not recipient or not msg_text:
                    result = "Xato: qabul qiluvchi yoki xabar matni bo'sh."
                elif BOT is None:
                    result = "Xato: bot hali tayyor emas, birozdan keyin urinib ko'ring."
                else:
                    matches = await asyncio.to_thread(db_find_business_chats, user_id, recipient)
                    if not matches:
                        known = await asyncio.to_thread(db_list_business_chats, user_id)
                        if known:
                            names = ", ".join(
                                f"{m['name']}" + (f" (@{m['username']})" if m["username"] else "")
                                for m in known
                            )
                            result = (
                                f"'{recipient}' topilmadi. Mavjud suhbatdoshlar: {names}. "
                                "Shulardan birini tanlang."
                            )
                        else:
                            result = (
                                f"'{recipient}' topilmadi. Hozircha ro'yxat bo'sh — suhbatdosh sizga "
                                "Telegram Business orqali kamida bitta xabar yozgan bo'lishi kerak, "
                                "shundan keyin unga xabar yubora olaman."
                            )
                    elif len(matches) > 1 and not any(
                        (m["name"] or "").lower() == recipient.lower()
                        or (m["username"] or "").lower() == recipient.lstrip("@").lower()
                        for m in matches
                    ):
                        names = ", ".join(m["name"] or m["username"] or "?" for m in matches)
                        result = f"Bir nechta mos suhbatdosh topildi: {names}. Qaysi biriga yuborishni aniqlashtiring."
                    else:
                        target = next(
                            (m for m in matches if (m["name"] or "").lower() == recipient.lower()
                             or (m["username"] or "").lower() == recipient.lstrip("@").lower()),
                            matches[0],
                        )
                        try:
                            await BOT.send_message(
                                chat_id=target["chat_id"], text=msg_text[:4000],
                                business_connection_id=target["conn_id"],
                            )
                            result = f"✅ Xabar {target['name'] or recipient} ga muvaffaqiyatli yuborildi."
                        except Exception as e:
                            logger.exception("Biznes xabar yuborishda xato")
                            result = f"Xabar yuborib bo'lmadi: {e}"
            elif fc.name == "generate_image":
                prompt = args.get("prompt", "")
                img = await asyncio.to_thread(do_generate_image, prompt)
                if img is not None and image_sink is not None:
                    image_sink.append((prompt, img))
                    result = "Rasm muvaffaqiyatli yaratildi va foydalanuvchiga yuborildi."
                else:
                    result = "Rasm yaratib bo'lmadi (xizmat vaqtincha ishlamayapti)."
            elif fc.name == "open_app" and device_action_sink is not None:
                app_name = args.get("app_name", "")
                device_action_sink.append({"type": "open_app", "app_name": app_name})
                result = f"'{app_name}' ilovasi ochilmoqda."
            elif fc.name == "set_alarm" and device_action_sink is not None:
                hour, minute = args.get("hour", 0), args.get("minute", 0)
                device_action_sink.append({
                    "type": "set_alarm", "hour": hour, "minute": minute,
                    "label": args.get("label", ""),
                    "date": args.get("date", ""),
                })
                result = f"Budilnik {hour:02d}:{minute:02d} ga o'rnatilmoqda."
            elif fc.name == "make_call" and device_action_sink is not None:
                phone = args.get("phone_number", "")
                contact = args.get("contact_name", "")
                device_action_sink.append({"type": "make_call", "phone_number": phone, "contact_name": contact})
                result = f"{contact or phone} ga qo'ng'iroq qilinmoqda."
            elif fc.name == "search_in_app" and device_action_sink is not None:
                app_name = args.get("app_name", "")
                query = args.get("query", "")
                video_url = ""
                if "youtube" in app_name.lower():
                    video_url = await asyncio.to_thread(resolve_youtube_video, query)
                device_action_sink.append({
                    "type": "search_in_app", "app_name": app_name, "query": query, "video_url": video_url,
                })
                if video_url:
                    result = f"'{query}' YouTube'da topildi va ijro etilmoqda."
                else:
                    result = f"{app_name} ilovasida '{query}' qidirilmoqda."
            elif fc.name == "open_telegram_chat" and device_action_sink is not None:
                username = args.get("username", "")
                msg = args.get("message", "")
                device_action_sink.append({"type": "open_telegram_chat", "username": username, "message": msg})
                result = f"Telegram'da {username} bilan suhbat ochilmoqda."
            elif fc.name == "send_sms" and device_action_sink is not None:
                phone = args.get("phone_number", "")
                msg = args.get("message", "")
                device_action_sink.append({"type": "send_sms", "phone_number": phone, "message": msg})
                result = f"{phone} raqamiga SMS tayyorlanmoqda."
            elif fc.name == "set_volume" and device_action_sink is not None:
                direction = args.get("direction", "up")
                device_action_sink.append({"type": "set_volume", "direction": direction})
                result = "Ovoz sozlanmoqda."
            elif fc.name == "toggle_flashlight" and device_action_sink is not None:
                on = bool(args.get("on", True))
                device_action_sink.append({"type": "toggle_flashlight", "on": on})
                result = "Fonar yoqilmoqda." if on else "Fonar o'chirilmoqda."
            else:
                result = await asyncio.to_thread(execute_function, user_id, fc.name, args)
            result_parts.append(types.Part.from_function_response(name=fc.name, response={"result": result}))
        contents.append(types.Content(role="user", parts=result_parts))

    # Har qanday holatda ham kontekstni saqlaymiz
    await asyncio.to_thread(_save_history, user_id, user_parts, answer, chat_id)
    return answer


_CYR2LAT = str.maketrans({
    "а": "a", "б": "b", "в": "v", "г": "g", "д": "d", "е": "e", "ё": "yo", "ж": "j",
    "з": "z", "и": "i", "й": "y", "к": "k", "л": "l", "м": "m", "н": "n", "о": "o",
    "п": "p", "р": "r", "с": "s", "т": "t", "у": "u", "ф": "f", "х": "x", "ц": "ts",
    "ч": "ch", "ш": "sh", "щ": "sh", "ъ": "", "ь": "", "э": "e", "ю": "yu", "я": "ya",
    "ў": "o'", "қ": "q", "ғ": "g'", "ҳ": "h",
})


def _detect_reply_override(text: str) -> str | None:
    """Foydalanuvchi shu xabarda javob turini so'ragan bo'lsa — aniqlaymiz.
    Masalan: "matnda javob ber" -> text, "ovozli xabarda javob ber" -> voice.
    Kirillcha yozilgan bo'lsa ham tushunadi."""
    t = text.lower().translate(_CYR2LAT)
    text_kw = (
        "matnda javob", "matn bilan javob", "matnda ber", "matnda yoz",
        "yozib javob", "yozma javob", "matnli javob", "tekstda javob", "matnda ayt",
        "matn tarzida", "tekst tarzida", "matn shaklida", "matn ko'rinishida",
        "matnda qaytar", "yozuvda javob",
    )
    voice_kw = (
        "ovozli javob", "ovozda javob", "ovozli xabarda", "ovoz bilan javob",
        "golosda javob", "golos bilan javob", "audio javob", "ovozli qilib", "ovozda ayt",
        "ovozli xabar tarzida", "ovozli xabar bilan", "ovozli xabar qilib",
        "ovoz tarzida", "ovoz shaklida", "ovozda qaytar", "ovozda ber", "golosda ayt",
    )
    if any(k in t for k in text_kw):
        return "text"
    if any(k in t for k in voice_kw):
        return "voice"
    return None


async def _reply_with_voice(message: Message, answer: str) -> bool:
    """Javobni ovozga aylantirib yuboradi. Muvaffaqiyatli bo'lsa True."""
    tts_text = answer[:1500]  # juda uzun matnni ovozga o'girish sekin va qimmat
    wav = await asyncio.to_thread(do_tts, tts_text)
    if not wav:
        return False
    ogg = await asyncio.to_thread(_wav_to_ogg, wav)
    try:
        if ogg:
            await message.answer_voice(BufferedInputFile(ogg, "javob.ogg"))
        else:
            await message.answer_audio(BufferedInputFile(wav, "javob.wav"), title="Javob")
        return True
    except Exception:
        logger.exception("Ovozli javob yuborishda xato")
        return False


async def agent_respond(message: Message, uid: int, parts: list[types.Part]):
    """Agentdan javob olib, foydalanuvchi tanloviga qarab matn yoki ovozda yuboradi."""
    images: list = []
    answer = await ask_agent(uid, parts, images)
    for cap, img in images:
        try:
            await message.answer_photo(
                BufferedInputFile(img, "rasm.png"),
                caption=(cap[:1000] if cap else None),
            )
        except Exception:
            logger.exception("Rasm yuborishda xato")
    if not answer:
        return

    # Javob turi: shu xabardagi buyruq > profil sozlamasi
    user_text = " ".join(p.text for p in parts if p.text)
    mode = _detect_reply_override(user_text)
    if mode is None:
        mode = await asyncio.to_thread(db_get_reply_mode, uid)

    if mode == "voice":
        sent = await _reply_with_voice(message, answer)
        if sent:
            if len(answer) > 1500:
                # Ovozga sig'magan qismi yo'qolmasligi uchun to'liq matnni ham yuboramiz
                await send_long(message, answer)
            return
        # TTS ishlamasa — matnga tushamiz

    await send_long(message, answer)


# ============================================================
# TELEGRAM HANDLERLAR
# ============================================================

router = Router()


def admin_keyboard() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📊 Statistika", callback_data="admin_stats"),
         InlineKeyboardButton(text="👥 Foydalanuvchilar", callback_data="admin_users")],
        [InlineKeyboardButton(text="⏳ Kutayotganlar", callback_data="admin_pending"),
         InlineKeyboardButton(text="📢 Broadcast", callback_data="admin_broadcast")],
        [InlineKeyboardButton(text="🔴 Botni o'chir" if bot_enabled else "🟢 Botni yoq", callback_data="admin_toggle")],
    ])


@router.message(Command("admin"))
async def cmd_admin(message: Message):
    if message.from_user.id != ADMIN_ID:
        return
    await message.answer("🛠 Admin panel:", reply_markup=admin_keyboard())


@router.callback_query(F.data.startswith("admin_"))
async def admin_callback(callback: CallbackQuery):
    if callback.from_user.id != ADMIN_ID:
        await callback.answer("Ruxsat yo'q!")
        return

    global bot_enabled
    action = callback.data

    if action == "admin_stats":
        text = await asyncio.to_thread(db_admin_stats)
        await callback.message.edit_text(text, reply_markup=admin_keyboard())

    elif action == "admin_users":
        text = await asyncio.to_thread(db_admin_users)
        await callback.message.edit_text(text, reply_markup=admin_keyboard())

    elif action == "admin_pending":
        rows = await asyncio.to_thread(db_pending_users)
        if not rows:
            await callback.message.edit_text("Kutayotgan foydalanuvchilar yo'q.", reply_markup=admin_keyboard())
        else:
            buttons = []
            for r in rows:
                name = f"@{r['username']}" if r['username'] else r['full_name']
                buttons.append([
                    InlineKeyboardButton(text=f"✅ {name}", callback_data=f"approve_{r['user_id']}"),
                    InlineKeyboardButton(text="❌", callback_data=f"revoke_{r['user_id']}"),
                ])
            buttons.append([InlineKeyboardButton(text="⬅️ Orqaga", callback_data="admin_back")])
            await callback.message.edit_text(
                "⏳ Ruxsat kutayotganlar:",
                reply_markup=InlineKeyboardMarkup(inline_keyboard=buttons)
            )

    elif action == "admin_toggle":
        bot_enabled = not bot_enabled
        status = "🟢 Bot yoqildi!" if bot_enabled else "🔴 Bot o'chirildi!"
        await callback.message.edit_text(f"Admin panel:\n{status}", reply_markup=admin_keyboard())

    elif action == "admin_broadcast":
        await callback.message.edit_text(
            "📢 Broadcast xabar yuboring:\n/broadcast <xabar matni>",
            reply_markup=admin_keyboard()
        )

    elif action == "admin_back":
        await callback.message.edit_text("🛠 Admin panel:", reply_markup=admin_keyboard())

    elif action.startswith("approve_"):
        uid = int(action.split("_")[1])
        ok = await asyncio.to_thread(db_approve_user, uid)
        if ok:
            try:
                await start_onboarding(callback.bot, uid)
            except Exception:
                pass
        await callback.answer("✅ Ruxsat berildi!" if ok else "Foydalanuvchi topilmadi")
        rows = await asyncio.to_thread(db_pending_users)
        if not rows:
            await callback.message.edit_text("Kutayotgan foydalanuvchilar yo'q.", reply_markup=admin_keyboard())
        else:
            buttons = []
            for r in rows:
                name = f"@{r['username']}" if r['username'] else r['full_name']
                buttons.append([
                    InlineKeyboardButton(text=f"✅ {name}", callback_data=f"approve_{r['user_id']}"),
                    InlineKeyboardButton(text="❌", callback_data=f"revoke_{r['user_id']}"),
                ])
            buttons.append([InlineKeyboardButton(text="⬅️ Orqaga", callback_data="admin_back")])
            await callback.message.edit_text("⏳ Ruxsat kutayotganlar:", reply_markup=InlineKeyboardMarkup(inline_keyboard=buttons))
        return

    elif action.startswith("revoke_"):
        uid = int(action.split("_")[1])
        await asyncio.to_thread(db_revoke_user, uid)
        await callback.answer("❌ Rad etildi")

    await callback.answer()


async def start_onboarding(bot: Bot, user_id: int):
    onboarding_state[user_id] = {"step": "name"}
    await bot.send_message(
        user_id,
        "✅ Botdan foydalanishga ruxsat berildi!\n\n"
        "Salom! 👋 Men sizning shaxsiy *Sirdosh AI* agentingizman.\n\n"
        "Keling, siz bilan tanishamiz — men sizni yaxshiroq bilsam, ko'proq yordam bera olaman! 😊\n\n"
        "Avvalo, *ismingiz* nima?",
        parse_mode="Markdown"
    )


@router.message(Command("approve"))
async def cmd_approve(message: Message, bot: Bot):
    if message.from_user.id != ADMIN_ID:
        return
    parts = message.text.split()
    if len(parts) < 2:
        await message.answer("Foydalanish: /approve <user_id>")
        return
    uid = int(parts[1])
    ok = await asyncio.to_thread(db_approve_user, uid)
    if ok:
        try:
            await start_onboarding(bot, uid)
        except Exception:
            pass
        await message.answer(f"✅ {uid} ruxsat berildi.")
    else:
        await message.answer("Foydalanuvchi topilmadi.")


@router.message(Command("revoke"))
async def cmd_revoke(message: Message):
    if message.from_user.id != ADMIN_ID:
        return
    parts = message.text.split()
    if len(parts) < 2:
        await message.answer("Foydalanish: /revoke <user_id>")
        return
    uid = int(parts[1])
    await asyncio.to_thread(db_revoke_user, uid)
    await message.answer(f"❌ {uid} ruxsati olib tashlandi.")


@router.message(Command("broadcast"))
async def cmd_broadcast(message: Message, bot: Bot):
    if message.from_user.id != ADMIN_ID:
        return
    text = message.text.removeprefix("/broadcast").strip()
    if not text:
        await message.answer("Foydalanish: /broadcast <xabar matni>")
        return
    user_ids = await asyncio.to_thread(db_get_all_user_ids)
    sent, failed = 0, 0
    for uid in user_ids:
        try:
            await bot.send_message(uid, f"📢 {text}")
            sent += 1
        except Exception:
            failed += 1
    await message.answer(f"Broadcast tugadi: {sent} ta yuborildi, {failed} ta xato.")


@router.message(CommandStart())
async def cmd_start(message: Message):
    uid = message.from_user.id
    await asyncio.to_thread(db_track_user, uid, message.from_user.username, message.from_user.full_name)

    if not await asyncio.to_thread(db_is_approved, uid):
        # Adminga xabar yuboramiz
        if ADMIN_ID and BOT:
            name = f"@{message.from_user.username}" if message.from_user.username else message.from_user.full_name
            try:
                await BOT.send_message(
                    ADMIN_ID,
                    f"🔔 Yangi foydalanuvchi botdan foydalanmoqchi:\n"
                    f"👤 {name} (ID: {uid})\n\n"
                    f"Ruxsat berish uchun: /approve {uid}\nRad etish: /revoke {uid}"
                )
            except Exception:
                pass
        await message.answer(
            "Salom! 👋\n\n"
            "Botdan foydalanish uchun admin ruxsati kerak.\n"
            "So'rovingiz adminga yuborildi — tez orada javob olasiz! ⏳"
        )
        return

    await message.answer(
        "Salom! 👋 Men sizning shaxsiy yordamchingizman.\n\n"
        "🔎 Internetdan ma'lumot — \"Dollar kursi qancha?\"\n"
        "💰 Buxgalter — \"50 ming taksiga ketdi\" yoki chek RASMINI yuboring\n"
        "📊 Hisobot — \"Bu oy qancha sarfladim?\"\n"
        "⏰ Eslatma — \"Ertaga 9 da dorini eslatib qo'y\"\n"
        "📝 Qayd — \"Eslab qol: mashina raqami 01A777BB\"\n"
        "🎤 Hammasi golosda ham ishlaydi!\n\n"
        "Komandalar: /hisobot, /eslatmalar, /clear, /forget"
    )


@router.message(Command("hisobot"))
async def cmd_report(message: Message):
    await message.answer(await asyncio.to_thread(db_get_report, message.from_user.id, "oy"))


@router.message(Command("eslatmalar"))
async def cmd_reminders(message: Message):
    await message.answer(await asyncio.to_thread(db_list_reminders, message.from_user.id))


@router.message(Command("clear"))
async def cmd_clear(message: Message):
    chat_history.pop((message.from_user.id, 0), None)
    await asyncio.to_thread(db_clear_history, message.from_user.id, 0)
    await message.answer("Suhbat tarixi tozalandi ✅")


# ============================================================
# LOYIHA MASLAHATCHISI — biznes g'oya → strategiya → UX → dizayn promptlari
# ============================================================

PROJECT_CHAT_ID = 777  # alohida kontekst: kundalik suhbat bilan aralashmaydi
project_mode_users: set[int] = set()  # hozir loyiha rejimida bo'lganlar

_PROJECT_STAGES = [
    ("g'oya", "Biznes g'oya va qiymat taklifi"),
    ("auditoriya", "Maqsadli auditoriya va muammolar"),
    ("model", "Daromad modeli va bozor"),
    ("raqobat", "Raqobat va noyob ustunlik"),
    ("funksiya", "Asosiy funksiyalar (MVP)"),
    ("ux", "Foydalanuvchi sayohati va onboarding"),
    ("kpi", "Muvaffaqiyat ko'rsatkichlari va strategiya"),
    ("hujjat", "Yakuniy hujjat + UI/UX dizayn promptlari"),
]


def build_project_prompt(user_id: int) -> str:
    profile = db_get_profile(user_id)
    name = (profile or {}).get("name") or "do'stim"
    stages = "\n".join(f"  {i+1}. {t}" for i, (_, t) in enumerate(_PROJECT_STAGES))
    return f"""Sen SIRDOSH — {name} uchun shaxsiy MAHSULOT STRATEGI, BIZNES-TAHLILCHI va UX-ARXITEKTORSAN.
{name} — veb-dasturchi. U yangi raqamli platforma/veb-ilova g'oyasi bilan keldi. Sening vazifang —
uni g'oyadan to dizaynerga beriladigan tayyor UI/UX promptlarigacha PROFESSIONAL darajada olib borish.

ISHLASH USULI — bosqichma-bosqich, dialog tarzida:
{stages}

HAR BOSQICHDA:
- 1-3 ta o'tkir, strategik savol ber (hammasini birdan emas!). Har savol NIMA UCHUN muhimligini bir jumlada ayt.
- Javobni TAHLIL qil: kuchli tomonini ta'kidla, zaif joyini ochiq ayt, "bunday qilsak yaxshiroq" deb taklif ber.
- Foydalanuvchi bilmasa — o'zing 2-3 ta professional variant taklif qil, u tanlasin.
- Bosqich yetarli ochilganda: "✅ [Bosqich] tugadi. Xulosa: ..." deb 2-3 jumla yakunlab, keyingisiga o't.
- Xabar boshida qaysi bosqichdaligingni ko'rsat, masalan: "📍 3/8 — Daromad modeli".

SAVOL NAMUNALARI (moslab ishlat, so'zma-so'z emas):
- Asosiy g'oya nima, qanday qiymat taklifi beradi? Foydalanuvchi qanday "og'riq"dan qutuladi?
- Kim uchun? Ularning hozirgi yechimi nima va nega u yetarli emas?
- Daromad: obuna / komissiya / reklama / sotuv / freemium — qaysi biri va nega?
- Raqobatchilar kim, siz nimada 10x yaxshisiz? Noyob savdo taklifi (USP)?
- KPI: qaysi 3 ta raqam muvaffaqiyatni ko'rsatadi (retention, LTV, konversiya...)?
- MVP: eng muhim 3-5 funksiya? Nima 2-versiyaga qoladi?
- Onboarding: birinchi 60 soniyada foydalanuvchi qanday "aha!" moment oladi?
- 1 yillik strategiya: qayerga boradi, qanday o'sadi?

YAKUNIY HUJJAT (8-bosqich) — quyidagi tuzilmada, to'liq va batafsil:
1. Loyiha nomi va bir jumlalik pitch
2. Muammo → Yechim → Qiymat taklifi
3. Maqsadli auditoriya: 2-3 ta persona (ism, yosh, kasb, maqsad, og'riq)
4. Biznes model va daromad manbalari
5. Raqobat tahlili va USP
6. MVP funksiyalar (ustuvorlik bilan) + keyingi versiyalar
7. Foydalanuvchi sayohati (user journey) — asosiy oqim qadam-baqadam
8. Axborot arxitekturasi — sahifalar/ekranlar daraxti
9. KPI va 12 oylik strategiya
10. UI/UX DIZAYN PROMPTLARI — har asosiy ekran uchun ALOHIDA, dizaynerga to'g'ridan-to'g'ri berish mumkin bo'lgan darajada:
    - Ekran maqsadi va foydalanuvchi bu yerda nima qiladi
    - Layout va tarkib ierarxiyasi
    - Vizual uslub: rang palitrasi (hex bilan), tipografika, ikonografiya, bo'shliqlar
    - Navigatsiya va foydalanuvchi oqimlari
    - Interaktivlik: mikro-interaksiyalar, hover/tap holatlari, animatsiya tavsiyalari
    - Bo'sh/yuklanish/xato holatlari
    - Mobil va desktop farqlari
    Promptlarni ingliz tilida ham ber (dizayn AI vositalari — Figma AI, Midjourney, v0, Lovable uchun).

QOIDALAR:
- Professional, lekin do'stona. Umumiy gaplar EMAS — aniq, amaliy, misollar bilan.
- Zaif g'oyaga xushomad qilma — halol tahlil ber va kuchaytirish yo'lini ko'rsat.
- Uzun matnlarni sarlavha va ro'yxatlar bilan tuzilmali yoz.
- Foydalanuvchi "hujjatni yoz", "yakunla", "promptlarni ber" desa — mavjud ma'lumot bilan darhol yakuniy hujjatni yoz (yetishmagan joyni o'zing eng oqilona taxmin bilan to'ldirib, buni belgilab qo'y).
- Foydalanuvchi "chiqish", "tugatdik" desa — qisqa xulosa ber va /loyiha_chiqish ni eslat.
- Til: foydalanuvchi qaysi tilda yozsa — o'sha (asosan o'zbek).

HOZIRGI VAQT: {now_local().strftime('%Y-%m-%d %H:%M')}."""


@router.message(Command("loyiha"))
async def cmd_project(message: Message):
    uid = message.from_user.id
    if not await asyncio.to_thread(db_is_approved, uid):
        await message.answer("Botdan foydalanish uchun admin ruxsati kerak. /start bosing.")
        return
    project_mode_users.add(uid)
    text = message.text.removeprefix("/loyiha").strip()
    intro = (
        "🚀 LOYIHA MASLAHATCHISI rejimi yoqildi\n\n"
        "Endi men sizning mahsulot strategingiz va UX-arxitektoringizman. "
        "G'oyangizni bosqichma-bosqich tahlil qilamiz:\n"
        "g'oya → auditoriya → daromad → raqobat → MVP → UX → KPI → "
        "yakuniy hujjat + dizayn promptlari.\n\n"
        "Bu suhbat alohida saqlanadi — kundalik ishlaringizga aralashmaydi. "
        "Istalgan payt qaytib davom ettirishingiz mumkin.\n"
        "Chiqish: /loyiha_chiqish\n\n"
    )
    if text:
        await message.answer(intro + "G'oyangizni o'qiyapman...")
        await agent_respond_project(message, uid, text)
    else:
        await message.answer(intro + "Boshlaymiz — g'oyangizni erkin, o'z so'zlaringiz bilan yozib bering. Qanday platforma va kim uchun?")


@router.message(Command("loyiha_chiqish"))
async def cmd_project_exit(message: Message):
    project_mode_users.discard(message.from_user.id)
    await message.answer(
        "Loyiha rejimidan chiqdik 👍 Oddiy suhbatga qaytdik.\n"
        "Loyihaga qaytish uchun yana /loyiha yozing — hamma narsa eslab qolingan."
    )


async def agent_respond_project(message: Message, uid: int, text: str):
    """Loyiha rejimida javob: alohida rol + alohida kontekst, javob har doim matnda."""
    answer = await ask_agent(
        uid, [types.Part.from_text(text=text)],
        chat_id=PROJECT_CHAT_ID,
        system_prompt=build_project_prompt(uid),
        tools_override=SAFE_BUSINESS_DECLARATIONS + [
            d for d in FUNCTION_DECLARATIONS
            if d.name in ("remember", "add_note", "find_notes", "get_note", "list_tags")
        ],
    )
    if answer:
        await send_long(message, answer)


def settings_keyboard(mode: str) -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[[
        InlineKeyboardButton(
            text=("✅ " if mode == "text" else "") + "📝 Matnli javob",
            callback_data="mode_text",
        ),
        InlineKeyboardButton(
            text=("✅ " if mode == "voice" else "") + "🔊 Ovozli javob",
            callback_data="mode_voice",
        ),
    ]])


@router.message(Command("sozlamalar"))
async def cmd_settings(message: Message):
    mode = await asyncio.to_thread(db_get_reply_mode, message.from_user.id)
    await message.answer(
        "⚙️ Sozlamalar\n\nMen javoblarni qanday beray?\n"
        "(Istalgan payt xabaringizda \"matnda javob ber\" yoki \"ovozli javob ber\" "
        "desangiz — o'sha safar aytganingizcha qilaman.)",
        reply_markup=settings_keyboard(mode),
    )


@router.callback_query(F.data.startswith("mode_"))
async def mode_callback(callback: CallbackQuery):
    mode = "voice" if callback.data == "mode_voice" else "text"
    await asyncio.to_thread(db_set_reply_mode, callback.from_user.id, mode)
    label = "🔊 Endi ovozli javob beraman." if mode == "voice" else "📝 Endi matnli javob beraman."
    try:
        await callback.message.edit_text(
            f"⚙️ Sozlamalar\n\n{label}", reply_markup=settings_keyboard(mode),
        )
    except Exception:
        pass  # xuddi shu tugma qayta bosilsa Telegram "not modified" beradi
    await callback.answer("Saqlandi ✅")


@router.message(Command("forget"))
async def cmd_forget(message: Message):
    kb = InlineKeyboardMarkup(inline_keyboard=[[
        InlineKeyboardButton(text="Ha, xotirani o'chir", callback_data="forget_yes"),
        InlineKeyboardButton(text="Yo'q", callback_data="forget_no"),
    ]])
    await message.answer(
        "⚠️ Bu men siz haqingizda eslab qolgan barcha narsalarni o'chiradi "
        "(loyihalaringiz, maqsadlaringiz...). Profil ma'lumotlari saqlanadi.\n\nRostdan o'chiraymi?",
        reply_markup=kb,
    )


@router.callback_query(F.data.startswith("forget_"))
async def forget_callback(callback: CallbackQuery):
    if callback.data == "forget_yes":
        uid = callback.from_user.id
        for k in [k for k in chat_history if k[0] == uid]:
            chat_history.pop(k, None)
        await asyncio.to_thread(db_clear_history, uid)
        n = await asyncio.to_thread(db_clear_memory, uid)
        await callback.message.edit_text(f"Xotira tozalandi ✅ ({n} ta yozuv o'chirildi).")
    else:
        await callback.message.edit_text("Bekor qilindi. Hech narsa o'chirilmadi 👍")
    await callback.answer()


# ============================================================
# TELEGRAM BUSINESS — egasi nomidan avtomatik javob
# ============================================================

_business_owners: dict[str, int] = {}  # connection_id -> egasining user_id si


async def _business_owner_id(bot: Bot, conn_id: str) -> int | None:
    if conn_id in _business_owners:
        return _business_owners[conn_id]
    try:
        conn = await bot.get_business_connection(conn_id)
        _business_owners[conn_id] = conn.user.id
        return conn.user.id
    except Exception:
        logger.exception("Business connection ma'lumotini olishda xato")
        return None


business_setup_state: dict[int, str] = {}  # uid -> "info" | "hours"


async def send_business_summary_to(owner_id: int, on_demand: bool = False) -> bool:
    """Bugungi biznes xulosani egasiga yuboradi. Xabar bo'lmasa faqat on_demand'da aytadi."""
    rows = await asyncio.to_thread(db_business_today, owner_id)
    if not rows:
        if on_demand and BOT:
            await BOT.send_message(owner_id, "📊 Bugun hech kim yozmagan.")
        return False

    per: dict[str, int] = {}
    for r in rows:
        s = r["sender"] or "Noma'lum"
        per[s] = per.get(s, 0) + 1
    top = sorted(per.items(), key=lambda x: -x[1])[:10]
    lines = "\n".join(f"• {s}: {c} ta xabar" for s, c in top)

    # AI mavzu xulosasi (buyurtma/shikoyat/narx so'rovlarini ajratib beradi)
    digest = ""
    try:
        msgs = "\n".join(f"- {r['sender']}: {r['text']}" for r in rows[:50])
        resp = await asyncio.to_thread(
            gemini_generate, model=MODEL,
            contents=(
                "Quyida bugun biznes egasiga mijozlardan kelgan xabarlar. 2-3 jumlada asosiy "
                "mavzularni xulosala — ayniqsa buyurtma, shikoyat, narx so'rovi bo'lsa alohida ayt. "
                f"O'zbekcha, qisqa:\n{msgs}"
            ),
            config=types.GenerateContentConfig(
                temperature=0.3, max_output_tokens=512,
                thinking_config=types.ThinkingConfig(thinking_budget=0),
            ),
        )
        if resp.candidates:
            try:
                digest = (resp.text or "").strip()
            except Exception:
                digest = ""
    except Exception:
        logger.exception("Biznes xulosa (AI) xato")

    text = (
        "📊 BUGUNGI BIZNES XULOSA\n\n"
        f"👥 {len(per)} kishi yozdi, jami {len(rows)} ta xabar:\n{lines}"
        + (f"\n\n🧠 Xulosa: {digest}" if digest else "")
    )
    if BOT:
        await BOT.send_message(owner_id, text[:4000])
    return True


async def business_daily_job():
    """Har kuni 21:00 da avto-javob yoqiq egalarga kunlik xulosa."""
    try:
        owners = await asyncio.to_thread(db_business_owners_with_auto)
        for uid in owners:
            try:
                await send_business_summary_to(uid)
            except Exception:
                logger.exception("Kunlik biznes xulosa yuborishda xato (user %s)", uid)
        await asyncio.to_thread(db_business_log_cleanup)
    except Exception:
        logger.exception("Kunlik biznes job xato")


def business_keyboard(p: dict) -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(
            text=("🔴 Avto-javobni o'chirish" if p["business_auto"] else "🟢 Avto-javobni yoqish"),
            callback_data="biz_toggle",
        )],
        [InlineKeyboardButton(text="📋 Biznes ma'lumotlari", callback_data="biz_info"),
         InlineKeyboardButton(text="⏰ Ish vaqti", callback_data="biz_hours")],
        [InlineKeyboardButton(
            text=("🔔 Xabarnoma: yoniq" if p["business_notify"] else "🔕 Xabarnoma: o'chiq"),
            callback_data="biz_notify",
        )],
        [InlineKeyboardButton(text="📊 Bugungi statistika", callback_data="biz_stats")],
    ])


def business_panel_text(p: dict) -> str:
    status = "🟢 YONIQ" if p["business_auto"] else "🔴 O'CHIQ"
    info = p["business_info"] or "— kiritilmagan (tugmani bosib kiriting)"
    if len(info) > 200:
        info = info[:200] + "..."
    hours = p["business_hours"] or "— cheklanmagan (doim javob beradi)"
    return (
        "💼 BIZNES PANEL\n\n"
        f"Avto-javob: {status}\n"
        f"Ish vaqti: {hours}\n"
        f"Biznes ma'lumotlari: {info}\n\n"
        "Qanday ishlaydi: sizga yozgan odamlarga men sizning nomingizdan, "
        "biznes ma'lumotlaringiz asosida javob beraman. Har suhbatdosh bilan "
        "kontekstni alohida eslab qolaman.\n\n"
        "Ulash: Telegram → Sozlamalar → Telegram Business → Chatbotlar → shu bot.\n"
        "(Telegram Business uchun Premium kerak.)"
    )


@router.message(Command("biznes"))
async def cmd_business(message: Message):
    uid = message.from_user.id
    if not await asyncio.to_thread(db_is_approved, uid):
        await message.answer("Botdan foydalanish uchun admin ruxsati kerak. /start bosing.")
        return
    p = await asyncio.to_thread(db_get_business_profile, uid)
    await message.answer(business_panel_text(p), reply_markup=business_keyboard(p))


@router.callback_query(F.data.startswith("biz_"))
async def business_callback(callback: CallbackQuery):
    uid = callback.from_user.id
    action = callback.data

    if action == "biz_toggle" or action == "biz_notify":
        field = "business_auto" if action == "biz_toggle" else "business_notify"
        p = await asyncio.to_thread(db_get_business_profile, uid)
        new_value = not p[field]
        await asyncio.to_thread(db_set_business_field, uid, field, new_value)
        p[field] = new_value
        try:
            await callback.message.edit_text(business_panel_text(p), reply_markup=business_keyboard(p))
        except Exception:
            pass
        await callback.answer("Saqlandi ✅")

    elif action == "biz_info":
        business_setup_state[uid] = "info"
        await callback.message.answer(
            "📋 Biznesingiz haqida yozib yuboring — nima ish qilasiz, mahsulot/xizmatlar, "
            "narxlar, manzil, yetkazib berish... Qancha batafsil bo'lsa, men mijozlarga "
            "shuncha aniq javob beraman.\n\n"
            "Bekor qilish uchun: «bekor» deb yozing."
        )
        await callback.answer()

    elif action == "biz_hours":
        business_setup_state[uid] = "hours"
        await callback.message.answer(
            "⏰ Ish vaqtingizni yozing (masalan: 09:00-18:00, dam olish: yakshanba).\n"
            "Ish vaqtidan tashqari kelgan xabarlarga buni hisobga olib javob beraman.\n\n"
            "Bekor qilish uchun: «bekor» deb yozing."
        )
        await callback.answer()

    elif action == "biz_stats":
        await callback.answer()
        await send_business_summary_to(uid, on_demand=True)


@router.business_message(F.text)
async def handle_business_message(message: Message, bot: Bot):
    conn_id = message.business_connection_id
    if not conn_id:
        return
    owner_id = await _business_owner_id(bot, conn_id)
    if owner_id is None:
        return
    # Egasining o'zi yozgan xabari yoki bot xabari — javob bermaymiz
    if not message.from_user or message.from_user.id == owner_id or message.from_user.is_bot:
        return
    profile = await asyncio.to_thread(db_get_business_profile, owner_id)
    if not profile["business_auto"]:
        return
    if not await asyncio.to_thread(db_is_approved, owner_id):
        return
    try:
        sender = message.from_user.full_name or "Suhbatdosh"

        # Suhbatdoshni eslab qolamiz — keyin egasi "unga yozib yubor" deya oladi
        await asyncio.to_thread(
            db_upsert_business_chat, owner_id, message.chat.id, conn_id,
            sender, message.from_user.username,
        )
        # Kunlik statistika uchun jurnalga yozamiz
        await asyncio.to_thread(
            db_business_log_add, owner_id, message.chat.id, sender, message.text,
        )

        # Egasining ismi — javoblar uning tilidan yoziladi
        owner_profile = await asyncio.to_thread(db_get_profile, owner_id)
        owner_name = (owner_profile or {}).get("name") or "xo'jayin"

        info_part = (
            f"\n\nBIZNES MA'LUMOTLARI (narx/shartlarni FAQAT shu yerdan ol, o'zingdan to'qima):\n{profile['business_info']}"
            if profile["business_info"] else
            "\n\nBiznes ma'lumotlari kiritilmagan — aniq savollar (narx, shartlar) bo'lsa "
            "'keyinroq aniq aytaman' deb javob ber."
        )
        hours_part = ""
        if profile["business_hours"]:
            hours_part = (
                f"\nIsh vaqti: {profile['business_hours']}. Hozir: {now_local().strftime('%H:%M, %A')}. "
                "Ish vaqtidan tashqari bo'lsa — buni tabiiy aytib, ish vaqtida to'liq javob berishingni bildir."
            )

        system = (
            f"SEN {owner_name.upper()}SAN — Telegram'da o'z suhbatdoshlaringga o'zing javob beryapsan. "
            "Birinchi shaxsda, tabiiy, samimiy, xuddi oddiy insonday yoz.\n"
            "QAT'IY QOIDALAR:\n"
            f"- Sen AI, bot, agent yoki yordamchi EMASSAN — sen {owner_name}san. Bu so'zlarni og'zingga ham olma.\n"
            f"- 'Sen kimsan?' desa — '{owner_name}man' de, vassalom. 'AI agentman' deyish QAT'IYAN TAQIQLANADI.\n"
            "- DIQQAT: agar shu suhbat tarixida oldin o'zingni AI/agent deb aytgan bo'lsang — u XATO edi, "
            "endi hech qachon takrorlama va o'sha uslubga ergashma.\n"
            "- 'Sizga qanday yordam bera olaman?', 'Eshitaman', 'Savolingiz bormi?' kabi call-center "
            "iboralari TAQIQLANADI. Oddiy odam kabi: salomga qisqa salom, savolga to'g'ridan-to'g'ri javob.\n"
            "- Javoblar QISQA — 1-3 jumla, do'stona ohang.\n"
            "- Umumiy savollarga (kurs, ob-havo, fakt) bilganingcha yoki web_search bilan aniq javob ber — "
            "'keyinroq aytaman' deb qochma.\n"
            "- FAQAT biznes narx/shartlari haqida aniq ma'lumot bo'lmasa — 'buni keyinroq aniqlab aytaman' de.\n"
            "- Shikoyat bo'lsa: samimiy uzr so'ra, 'tez orada o'zim hal qilaman' de.\n"
            "- Suhbatdosh qaysi tilda yozsa, o'sha tilda javob ber."
            f"{info_part}{hours_part}\n\n"
            f"Hozirgi suhbatdosh: {sender}."
        )
        # Har suhbatdosh bilan alohida kontekst (chat_id = suhbat raqami).
        # Faqat xavfsiz funksiyalar — begona odam egasining moliya/eslatmalariga tega olmaydi.
        answer = await ask_agent(
            owner_id, [types.Part.from_text(text=message.text)], chat_id=message.chat.id,
            system_prompt=system, tools_override=SAFE_BUSINESS_DECLARATIONS,
        )
        if not answer:
            return
        await bot.send_message(
            chat_id=message.chat.id, text=answer[:4000],
            business_connection_id=conn_id,
        )
        # Egasiga xabarnoma: kim yozdi, nima deb javob berdim
        if profile["business_notify"]:
            try:
                uname = f" (@{message.from_user.username})" if message.from_user.username else ""
                await bot.send_message(
                    owner_id,
                    f"💼 {sender}{uname} sizga yozdi:\n«{message.text[:300]}»\n\n"
                    f"Men javob berdim:\n«{answer[:300]}»",
                )
            except Exception:
                logger.exception("Egasiga xabarnoma yuborishda xato")
    except Exception:
        logger.exception("Biznes avto-javobda xato")


_TRANSCRIBE_PROMPT = (
    "TRANSKRIPSIYA VAZIFASI. Sen tarjimon emassan, faqat transkriptchisan.\n"
    "Quyidagi audioda inson nima gapirgan bo'lsa, FAQAT o'sha so'zlarni, xuddi shu tilda yoz.\n"
    "QATIY TAQIQ: bu ko'rsatmani, izoh, sarlavha, tirnoq belgisi yoki boshqa hech qanday qo'shimcha matn yozma.\n"
    "Agar audio bo'sh/tushunarsiz bo'lsa — bo'sh javob qaytar."
)


def _clean_transcript(text: str) -> str:
    """Model ba'zan o'z ko'rsatmasini ham qaytarib yuborishi mumkin — shularni tozalaymiz."""
    leak_markers = ("TRANSKRIPSIYA VAZIFASI", "Bu audio xabarni tinglaysan", "QATIY TAQIQ", "so'zma-so'z")
    if any(m.lower() in text.lower() for m in leak_markers):
        # Ko'rsatma sızib chiqqan — eng oxirgi qatorni yoki bo'sh qaytaramiz (ishonchsiz natija)
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        lines = [l for l in lines if not any(m.lower() in l.lower() for m in leak_markers)]
        return lines[-1] if lines else ""
    return text.strip()


async def transcribe_audio(data: bytes, mime: str) -> str:
    """Gemini orqali audio ni matnga o'giradi."""
    response = await asyncio.to_thread(
        gemini_generate,
        model=MODEL,
        contents=[
            types.Part.from_text(text=_TRANSCRIBE_PROMPT),
            types.Part.from_bytes(data=data, mime_type=mime),
        ],
        config=types.GenerateContentConfig(
            temperature=0.0,
            max_output_tokens=2048,
            thinking_config=types.ThinkingConfig(thinking_budget=0),
        ),
    )
    if not response.candidates:
        return ""
    try:
        raw = (response.text or "").strip()
    except Exception:
        parts = (response.candidates[0].content.parts or []) if response.candidates[0].content else []
        raw = " ".join(p.text for p in parts if p.text).strip()
    return _clean_transcript(raw)


def _fetch_telegram_file(file_id: str) -> bytes:
    """Faylni Telegramdan to'g'ridan-to'g'ri requests bilan yuklaydi.

    aiogram sessiyasi (polling bilan bir xil ulanish) ba'zan get_file/download_file
    da 60s timeout beryapti. requests alohida ulanish ochib, buni chetlab o'tadi.
    """
    base = f"https://api.telegram.org/bot{BOT_TOKEN}"
    meta = requests.get(f"{base}/getFile", params={"file_id": file_id}, timeout=30)
    meta.raise_for_status()
    file_path = meta.json()["result"]["file_path"]
    resp = requests.get(
        f"https://api.telegram.org/file/bot{BOT_TOKEN}/{file_path}", timeout=60
    )
    resp.raise_for_status()
    return resp.content


@router.message(F.voice | F.audio)
async def handle_voice(message: Message, bot: Bot):
    uid = message.from_user.id
    if not bot_enabled and uid != ADMIN_ID:
        await message.answer("Bot vaqtincha o'chirilgan. Tez orada qaytamiz!")
        return
    if not await asyncio.to_thread(db_is_approved, uid):
        await message.answer("Botdan foydalanish uchun admin ruxsati kerak. /start bosing.")
        return
    await asyncio.to_thread(db_track_user, uid, message.from_user.username, message.from_user.full_name)
    await bot.send_chat_action(message.chat.id, ChatAction.TYPING)
    try:
        audio = message.voice or message.audio
        if audio.file_size and audio.file_size > 20 * 1024 * 1024:
            await message.answer("Audio juda katta (20 MB dan oshmasin).")
            return
        data = await asyncio.to_thread(_fetch_telegram_file, audio.file_id)
        mime = "audio/ogg" if message.voice else (audio.mime_type or "audio/mpeg")

        # Avval ovozni matnga o'giramiz
        text = await transcribe_audio(data, mime)
        if not text:
            await message.answer("Ovozni tushunib bo'lmadi 😕 Iltimos qaytadan yuboring.")
            return

        # Keyin matn sifatida agentga yuboramiz
        if uid in project_mode_users:
            await agent_respond_project(message, uid, text)
        else:
            await agent_respond(message, message.from_user.id, [types.Part.from_text(text=text)])
    except Exception:
        logger.exception("Golosli xabarda xato")
        await message.answer("Xatolik yuz berdi 😕 Qaytadan urinib ko'ring.")


@router.message(F.photo)
async def handle_photo(message: Message, bot: Bot):
    uid = message.from_user.id
    if not bot_enabled and uid != ADMIN_ID:
        await message.answer("Bot vaqtincha o'chirilgan. Tez orada qaytamiz!")
        return
    if not await asyncio.to_thread(db_is_approved, uid):
        await message.answer("Botdan foydalanish uchun admin ruxsati kerak. /start bosing.")
        return
    await asyncio.to_thread(db_track_user, uid, message.from_user.username, message.from_user.full_name)
    await bot.send_chat_action(message.chat.id, ChatAction.TYPING)
    try:
        photo = message.photo[-1]  # eng katta o'lcham
        img_data = await asyncio.to_thread(_fetch_telegram_file, photo.file_id)
        caption = (message.caption or "").strip()

        # Rasmni eslab qolamiz — keyin "buni fonini o'chir" desa ishlatamiz (15 daqiqa)
        last_user_image[uid] = (img_data, "image/jpeg", time.time())

        # Izoh tahrir buyrug'imi (savol emas)? -> rasmni tahrirlaymiz
        if caption and _is_edit_instruction(caption) and not _is_analysis_question(caption):
            await edit_and_send(message, img_data, caption, "image/jpeg")
            return

        # Aks holda: tahlil qilamiz (chek/matn/tushuntirish)
        instruction = (
            "Bu rasmni tahlil qil. Agar chek/kvitansiya/to'lov rasmi bo'lsa — "
            "summani va do'kon/xizmat nomini aniqlab add_transaction funksiyasini chaqir, "
            "keyin nimani yozganingni ayt. Boshqa rasm bo'lsa, shunchaki tushuntir."
        )
        if caption:
            instruction += f"\nFoydalanuvchi izohi: {caption}"
        parts = [
            types.Part.from_bytes(data=img_data, mime_type="image/jpeg"),
            types.Part.from_text(text=instruction),
        ]
        await agent_respond(message, message.from_user.id, parts)
        if not caption:
            await message.answer(
                "💡 Bu rasmni tahrirlashni xohlasangiz — nima o'zgartirishni yozing.\n"
                "Masalan: «fonini o'chir», «ko'ylakni qizil qil», «anime uslubida chiz», «sifatini yaxshila»."
            )
    except Exception:
        logger.exception("Rasmda xato")
        await message.answer("Rasmni o'qishda xatolik 😕 Qaytadan urinib ko'ring.")


async def edit_and_send(message: Message, img_bytes: bytes, prompt: str, mime: str):
    """Rasmni tahrirlab, natijani yuboradi. Muvaffaqiyatsiz bo'lsa xabar beradi."""
    await message.bot.send_chat_action(message.chat.id, ChatAction.UPLOAD_PHOTO)
    out_img, out_text, reason = await asyncio.to_thread(do_edit_image, img_bytes, prompt, mime)
    if out_img:
        # Tahrirlangan rasmni keyingi tahrir uchun ham eslab qolamiz (zanjir bo'lib tahrirlash)
        last_user_image[message.from_user.id] = (out_img, "image/png", time.time())
        await message.answer_photo(
            BufferedInputFile(out_img, "tahrirlangan.png"),
            caption=(out_text[:1000] if out_text else "Mana, tayyor ✅ Yana o'zgartirish kerak bo'lsa yozing."),
        )
    elif reason == "quota":
        await message.answer(
            "⚠️ Rasm tahrirlash limiti tugadi. Gemini'ning bepul tarifida rasm "
            "generatsiyasi kunlik cheklangan — ertaga tiklanadi yoki API kalitida billing yoqilsa cheksiz ishlaydi."
        )
    else:
        await message.answer(
            "Rasmni tahrirlab bo'lmadi 😕 Qaytadan, aniqroq yozib ko'ring "
            "(masalan: «orqa fonni oq qil»)."
        )


@router.message(F.document)
async def handle_document(message: Message, bot: Bot):
    uid = message.from_user.id
    if not bot_enabled and uid != ADMIN_ID:
        await message.answer("Bot vaqtincha o'chirilgan. Tez orada qaytamiz!")
        return
    if not await asyncio.to_thread(db_is_approved, uid):
        await message.answer("Botdan foydalanish uchun admin ruxsati kerak. /start bosing.")
        return
    await asyncio.to_thread(db_track_user, uid, message.from_user.username, message.from_user.full_name)
    await bot.send_chat_action(message.chat.id, ChatAction.TYPING)

    doc = message.document
    fname = (doc.file_name or "fayl").lower()
    caption = message.caption or ""

    if doc.file_size and doc.file_size > 20 * 1024 * 1024:
        await message.answer("Fayl juda katta (20 MB dan oshmasin).")
        return

    try:
        data = await asyncio.to_thread(_fetch_telegram_file, doc.file_id)

        instruction = caption or "Bu hujjatni tahlil qil, asosiy mazmunini va muhim nuqtalarini tushuntir."

        # PDF — Gemini to'g'ridan-to'g'ri o'qiydi
        if fname.endswith(".pdf") or doc.mime_type == "application/pdf":
            parts = [
                types.Part.from_bytes(data=data, mime_type="application/pdf"),
                types.Part.from_text(text=instruction),
            ]
            await agent_respond(message, uid, parts)
            return

        # Word / Excel / matn — matnni ajratamiz
        if fname.endswith(".docx"):
            text = await asyncio.to_thread(extract_docx, data)
        elif fname.endswith((".xlsx", ".xlsm")):
            text = await asyncio.to_thread(extract_xlsx, data)
        elif fname.endswith((".txt", ".csv", ".md", ".json")):
            text = data.decode("utf-8", errors="ignore")
        elif fname.endswith(".doc"):
            await message.answer("Eski .doc formati qo'llab-quvvatlanmaydi. Iltimos .docx ga aylantiring.")
            return
        else:
            await message.answer("Bu fayl turini o'qiy olmadim. PDF, Word (.docx), Excel (.xlsx) yoki matn yuboring.")
            return

        if not text.strip():
            await message.answer("Hujjatdan matn topilmadi (bo'sh yoki rasm ko'rinishida).")
            return

        text = text[:30000]  # juda katta hujjatlarni cheklash
        prompt = f"{instruction}\n\n=== HUJJAT MAZMUNI ({doc.file_name}) ===\n{text}"
        await agent_respond(message, uid, [types.Part.from_text(text=prompt)])
    except Exception:
        logger.exception("Hujjatda xato")
        await message.answer("Hujjatni o'qishda xatolik 😕 Qaytadan urinib ko'ring.")


@router.message(F.text)
async def handle_text(message: Message, bot: Bot):
    uid = message.from_user.id

    # Biznes panel sozlash jarayoni (/biznes -> ma'lumot yoki ish vaqti kiritish)
    if uid in business_setup_state:
        field = business_setup_state.pop(uid)
        text = message.text.strip()
        if text.lower() in ("bekor", "bekor qilish", "cancel", "отмена"):
            await message.answer("Bekor qilindi 👍")
            return
        if field == "info":
            await asyncio.to_thread(db_set_business_field, uid, "business_info", text[:3000])
            await message.answer(
                "📋 Biznes ma'lumotlari saqlandi! Endi mijozlarga shu asosda javob beraman.\n"
                "Panelga qaytish: /biznes"
            )
        else:
            await asyncio.to_thread(db_set_business_field, uid, "business_hours", text[:100])
            await message.answer("⏰ Ish vaqti saqlandi!\nPanelga qaytish: /biznes")
        return

    # Yaqinda rasm yuklagan bo'lsa: tahrir buyrug'i YOKI "rasm/surat/shunday qil" desa —
    # o'sha namuna rasmni ishlatib tahrirlaymiz (savol bo'lmasa). Namunani e'tiborsiz qoldirmaslik uchun.
    img_entry = last_user_image.get(uid)
    if img_entry and (time.time() - img_entry[2] < 900) and uid not in onboarding_state \
            and not _is_analysis_question(message.text) \
            and (_is_edit_instruction(message.text) or _mentions_image(message.text)):
        await edit_and_send(message, img_entry[0], message.text.strip(), img_entry[1])
        return

    # Onboarding jarayoni
    if uid in onboarding_state:
        state = onboarding_state[uid]
        text = message.text.strip()
        step = state["step"]

        if step == "name":
            state["name"] = text
            state["step"] = "profession"
            await message.answer(
                f"Juda yaxshi, *{text}*! 😊\n\n"
                "Siz qanday soha bilan shug'ullanasiz?\n"
                "_(Masalan: dasturchi, tadbirkor, talaba, shifokor...)_",
                parse_mode="Markdown"
            )

        elif step == "profession":
            state["profession"] = text
            state["step"] = "interests"
            await message.answer(
                "Zo'r! 💪\n\n"
                "Qiziqishlaringiz nima?\n"
                "_(Masalan: texnologiya, biznes, sport, musiqa, sayohat...)_",
                parse_mode="Markdown"
            )

        elif step == "interests":
            state["interests"] = text
            state["step"] = "goals"
            await message.answer(
                "Ajoyib! 🌟\n\n"
                "Men sizga eng ko'p qaysi sohada yordam bera olaman?\n"
                "_(Masalan: ish, o'qish, moliyaviy hisob, eslatmalar, ma'lumot qidirish...)_",
                parse_mode="Markdown"
            )

        elif step == "goals":
            state["goals"] = text
            await asyncio.to_thread(
                db_save_profile, uid,
                state["name"], state["profession"],
                state["interests"], state["goals"]
            )
            del onboarding_state[uid]
            await message.answer(
                f"Tanishganimizdan xursandman, *{state['name']}*! 🎉\n\n"
                f"Men endi siz haqingizda ko'proq bilaman va sizga samarali yordam bera olaman.\n\n"
                f"📌 Sizning profilingiz:\n"
                f"👤 Ism: {state['name']}\n"
                f"💼 Kasb: {state['profession']}\n"
                f"🎯 Qiziqishlar: {state['interests']}\n"
                f"🚀 Maqsadlar: {state['goals']}\n\n"
                f"Endi menga istalgan savolni bering — men doim yordamga tayyorman! 😊\n\n"
                f"Komandalar: /hisobot, /eslatmalar, /clear, /forget",
                parse_mode="Markdown"
            )
        return

    if not bot_enabled and uid != ADMIN_ID:
        await message.answer("Bot vaqtincha o'chirilgan. Tez orada qaytamiz!")
        return
    if not await asyncio.to_thread(db_is_approved, uid):
        await message.answer("Botdan foydalanish uchun admin ruxsati kerak. /start bosing.")
        return
    await asyncio.to_thread(db_track_user, uid, message.from_user.username, message.from_user.full_name)
    await bot.send_chat_action(message.chat.id, ChatAction.TYPING)
    try:
        if uid in project_mode_users:
            await agent_respond_project(message, uid, message.text)
        else:
            await agent_respond(message, message.from_user.id, [types.Part.from_text(text=message.text)])
    except Exception:
        logger.exception("Matnli xabarda xato")
        await message.answer("Xatolik yuz berdi 😕 Qaytadan urinib ko'ring.")


async def send_long(message: Message, text: str):
    """Uzun matnni bo'lib yuboradi. Kod bloki (```) bo'lsa Markdown bilan — Telegram'da
    "nusxalash" tugmasi chiqadi; Markdown xato bersa oddiy matnga tushadi."""
    for i in range(0, len(text), 4000):
        chunk = text[i:i + 4000]
        if "```" in chunk:
            try:
                await message.answer(chunk, parse_mode="Markdown")
                continue
            except Exception:
                pass  # noto'g'ri Markdown — oddiy matn sifatida yuboramiz
        await message.answer(chunk)


# ============================================================
# ADMIN OGOHLANTIRISHLARI (xato bo'lsa Telegram'da xabar keladi)
# ============================================================

async def _send_admin_alert(text: str):
    try:
        if BOT and ADMIN_ID:
            await BOT.send_message(ADMIN_ID, text)
    except Exception:
        pass  # ogohlantirish yuborilmasa ham bot ishlashda davom etadi


class AdminAlertHandler(logging.Handler):
    """logger.error/exception bo'lganda adminga Telegram xabar yuboradi.
    Bir xil xato 60 soniyada faqat 1 marta yuboriladi (flood bo'lmasligi uchun)."""

    def __init__(self, loop: asyncio.AbstractEventLoop):
        super().__init__(level=logging.ERROR)
        self.loop = loop
        self._last: dict[str, float] = {}

    def emit(self, record: logging.LogRecord):
        try:
            if not BOT or not ADMIN_ID:
                return
            key = f"{record.name}:{record.getMessage()}"
            now = time.time()
            if now - self._last.get(key, 0.0) < 60:
                return
            self._last[key] = now
            text = f"🚨 Xato: {record.getMessage()}"
            if record.exc_info and record.exc_info[1] is not None:
                text += f"\n{type(record.exc_info[1]).__name__}: {record.exc_info[1]}"
            asyncio.run_coroutine_threadsafe(_send_admin_alert(text[:1000]), self.loop)
        except Exception:
            pass


def _log_versions():
    """Ishlayotgan kutubxona versiyalarini logga yozamiz — keyin aniq pin qilish oson bo'ladi."""
    import importlib.metadata as md
    for pkg in ("aiogram", "google-genai", "psycopg2-binary", "fastapi", "APScheduler"):
        try:
            logger.info("Versiya: %s==%s", pkg, md.version(pkg))
        except Exception:
            pass


async def main():
    global BOT
    db_init()
    _log_versions()
    BOT = Bot(token=BOT_TOKEN)
    alert_handler = AdminAlertHandler(asyncio.get_running_loop())
    logger.addHandler(alert_handler)
    logging.getLogger("sirdosh-api").addHandler(alert_handler)
    dp = Dispatcher()
    dp.include_router(router)
    scheduler.start()
    restore_reminders()
    # Har kuni 21:00 da biznes egalariga kunlik xulosa
    scheduler.add_job(business_daily_job, "cron", hour=21, minute=0, id="biz_daily")
    try:
        await BOT.set_my_commands([
            BotCommand(command="start", description="Boshlash"),
            BotCommand(command="sozlamalar", description="⚙️ Javob turi: matn / ovoz"),
            BotCommand(command="biznes", description="💼 Avto-javob (Telegram Business)"),
            BotCommand(command="loyiha", description="🚀 Loyiha maslahatchisi: g'oya → strategiya → UX"),
            BotCommand(command="loyiha_chiqish", description="Loyiha rejimidan chiqish"),
            BotCommand(command="hisobot", description="Oylik moliyaviy hisobot"),
            BotCommand(command="eslatmalar", description="Faol eslatmalar"),
            BotCommand(command="clear", description="Suhbat tarixini tozalash"),
        ])
    except Exception:
        logger.exception("Komandalar menyusini o'rnatishda xato")
    logger.info("Shaxsiy yordamchi (v2) ishga tushdi...")
    await dp.start_polling(BOT)


if __name__ == "__main__":
    asyncio.run(main())
